from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse, HTMLResponse
from pydantic import BaseModel, Field
import difflib
from typing import List, Dict, Tuple
from collections import Counter
from datetime import datetime, timezone
from functools import lru_cache
import hashlib
import json
import os
import re
import unicodedata
import html as html_lib
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from urllib.parse import quote, urlencode
from qdrant_client import QdrantClient
from bs4 import BeautifulSoup
from docx import Document
from pdfminer.high_level import extract_text as pdf_extract_text

from config import JSON_DIR, QDRANT_PATH, COLLECTION_NAME, LEXICAL_INDEX_PATH
import retriever as retriever_module
from retriever import retrieve
from generator import generate, extract_sources, generate_local, generate_remote
from reference_replace import apply_reference_markers, search_dorar_candidates, search_local_hadith
from ingest_common import (
    infer_conversion_status,
    infer_document_type,
    infer_source_ext,
    infer_source_type,
    resolve_index_json_path,
    load_state,
    save_state,
    release_book,
)

app = FastAPI(title="RAG Buku Islam", version="0.1.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


def _load_index() -> Dict:
    index_path = os.path.join(JSON_DIR, "_index.json")
    with open(index_path, "r", encoding="utf-8") as f:
        return json.load(f)


def _index_path() -> str:
    return os.path.join(JSON_DIR, "_index.json")


def _content_index_path() -> str:
    return os.path.join(JSON_DIR, "_content_index.json")


def _load_json_file(path: str, default):
    if not os.path.exists(path):
        return default
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return default
    return data if isinstance(data, type(default)) else default


def _save_json_file(path: str, data) -> None:
    os.makedirs(os.path.dirname(path), exist_ok=True)
    tmp = f"{path}.tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2, sort_keys=True)
    os.replace(tmp, path)


def _book_entry(book_id: str) -> Dict | None:
    for record in _load_index().get("files", []):
        record_book_id = record.get("book_id") or os.path.splitext(record["filename"])[0]
        if record_book_id == book_id:
            return record
    return None


def _book_json_path(record: Dict) -> str:
    return resolve_index_json_path(record, JSON_DIR)


def _load_book_record(book_id: str) -> Dict | None:
    record = _book_entry(book_id)
    if not record:
        return None
    json_path = _book_json_path(record)
    if not os.path.exists(json_path):
        return None
    with open(json_path, "r", encoding="utf-8") as f:
        book = json.load(f)
    return {"record": record, "book": book, "json_path": json_path}


def _find_book_page(book: Dict, page_num: int) -> Dict | None:
    for p in book.get("pages", []):
        if p.get("page") == page_num:
            return p
    return None


def _build_unified_diff(before: str, after: str, from_label: str = "before", to_label: str = "after", limit: int = 9000) -> str:
    before_lines = (before or "").splitlines()
    after_lines = (after or "").splitlines()
    diff_lines = list(
        difflib.unified_diff(
            before_lines,
            after_lines,
            fromfile=from_label,
            tofile=to_label,
            lineterm="",
            n=3,
        )
    )
    if not diff_lines:
        return "Tidak ada perubahan."
    diff_text = "\n".join(diff_lines)
    if len(diff_text) > limit:
        diff_text = diff_text[:limit] + "\n... (diff dipotong)"
    return diff_text


def _normalize_signature_text(text: str) -> str:
    text = unicodedata.normalize("NFKC", text or "")
    text = text.replace("\x0c", " ").replace("\xa0", " ")
    text = text.lower()
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _sha256_text(text: str) -> str:
    return hashlib.sha256((text or "").encode("utf-8")).hexdigest()


def _simhash(tokens: List[str]) -> int:
    counts = Counter(tokens)
    if not counts:
        return 0
    vector = [0] * 64
    for token, weight in counts.items():
        digest = hashlib.blake2b(token.encode("utf-8"), digest_size=8).digest()
        fingerprint = int.from_bytes(digest, "big")
        for bit in range(64):
            if fingerprint & (1 << bit):
                vector[bit] += weight
            else:
                vector[bit] -= weight
    value = 0
    for bit, score in enumerate(vector):
        if score > 0:
            value |= 1 << bit
    return value


def _build_book_signature(book: Dict) -> Dict:
    pages = book.get("pages", []) or []
    page_hashes: List[str] = []
    normalized_pages: List[str] = []
    for page in pages:
        content = page.get("content", "")
        normalized = _normalize_signature_text(content)
        if not normalized:
            continue
        normalized_pages.append(normalized)
        page_hashes.append(_sha256_text(normalized))

    joined_text = "\n".join(normalized_pages)
    return {
        "page_hashes": page_hashes,
        "content_hash": _sha256_text("\n".join(page_hashes)),
        "text_hash": _sha256_text(joined_text),
        "text_simhash": _simhash(joined_text.split()),
        "page_count": len(page_hashes),
        "size_bytes": int(book.get("size_bytes", 0) or 0),
    }


def _find_index_record(index: Dict, book_id: str) -> Dict | None:
    for record in index.get("files", []):
        record_book_id = record.get("book_id") or os.path.splitext(record.get("filename", ""))[0]
        if record_book_id == book_id:
            return record
    return None


def _apply_book_to_record(record: Dict, book: Dict, json_path: str) -> Dict:
    updated = dict(record)
    updated["book_id"] = record.get("book_id") or os.path.splitext(record.get("filename", ""))[0]
    updated["filename"] = book.get("filename", updated.get("filename", ""))
    updated["language"] = book.get("language", updated.get("language", "unknown"))
    updated["title"] = book.get("title", updated.get("title", ""))
    updated["total_pages"] = int(book.get("total_pages", len(book.get("pages", []))) or len(book.get("pages", [])))
    updated["json_path"] = record.get("json_path", json_path.replace(JSON_DIR + os.sep, "").replace("\\", "/"))
    updated["source_root"] = record.get("source_root", "")
    updated["source_path"] = record.get("source_path", "")
    updated["source_relpath"] = record.get("source_relpath", "")
    updated["source_ext"] = record.get("source_ext", infer_source_ext(record))
    updated["source_type"] = record.get("source_type", infer_source_type(record))
    updated["document_type"] = record.get("document_type", infer_document_type(record))
    updated["conversion_status"] = record.get("conversion_status", infer_conversion_status(record))
    updated["quality_status"] = book.get("quality_status", record.get("quality_status", "ok"))
    updated["quality_reasons"] = book.get("quality_reasons", record.get("quality_reasons", []))
    updated["quality_warnings"] = book.get("quality_warnings", record.get("quality_warnings", []))
    updated["review_status"] = book.get("review_status", record.get("review_status", "approved_auto"))
    updated["review_required"] = bool(book.get("review_required", record.get("review_required", False)))
    updated["review_route"] = book.get("review_route", record.get("review_route", "auto"))
    updated["reviewed_by"] = book.get("reviewed_by", record.get("reviewed_by", ""))
    updated["reviewed_at"] = book.get("reviewed_at", record.get("reviewed_at", ""))
    updated["review_note"] = book.get("review_note", record.get("review_note", ""))
    updated["ingest_ready"] = bool(book.get("ingest_ready", record.get("ingest_ready", True)))
    return updated


def _apply_record_to_book_json(book: Dict, record: Dict, json_path: str) -> Dict:
    merged = dict(book)
    merged.setdefault("book_id", record.get("book_id") or os.path.splitext(record.get("filename", ""))[0])
    merged.setdefault("filename", record.get("filename", merged.get("filename", "")))
    merged.setdefault("json_path", record.get("json_path", os.path.relpath(json_path, JSON_DIR).replace("\\", "/")))
    merged.setdefault("source_root", record.get("source_root", ""))
    merged.setdefault("source_path", record.get("source_path", ""))
    merged.setdefault("source_relpath", record.get("source_relpath", ""))
    merged.setdefault("source_ext", record.get("source_ext", infer_source_ext(record)))
    merged.setdefault("source_type", record.get("source_type", infer_source_type(record)))
    merged.setdefault("document_type", record.get("document_type", infer_document_type(record)))
    merged.setdefault("conversion_status", record.get("conversion_status", infer_conversion_status(record)))
    merged.setdefault("source_hash", record.get("source_hash", ""))
    merged.setdefault("size_bytes", int(record.get("size_bytes", merged.get("size_bytes", 0)) or 0))
    merged.setdefault("quality_status", record.get("quality_status", "ok"))
    merged.setdefault("quality_reasons", record.get("quality_reasons", []))
    merged.setdefault("quality_warnings", record.get("quality_warnings", []))
    merged.setdefault("review_status", record.get("review_status", "approved_auto"))
    merged.setdefault("review_required", bool(record.get("review_required", False)))
    merged.setdefault("review_route", record.get("review_route", "auto"))
    merged.setdefault("reviewed_by", record.get("reviewed_by", ""))
    merged.setdefault("reviewed_at", record.get("reviewed_at", ""))
    merged.setdefault("review_note", record.get("review_note", ""))
    merged.setdefault("ingest_ready", bool(record.get("ingest_ready", True)))
    return merged


def _compact_page_nav(book: Dict, current_page: int, window: int = 4, edge: int = 3) -> List[int]:
    total_pages = int(book.get("total_pages", current_page) or current_page)
    pages = set()
    for p in range(1, min(edge, total_pages) + 1):
        pages.add(p)
    for p in range(max(1, current_page - window), min(total_pages, current_page + window) + 1):
        pages.add(p)
    for p in range(max(1, total_pages - edge + 1), total_pages + 1):
        pages.add(p)
    return sorted(pages)


def _highlight_terms_html(escaped_html: str, query: str) -> str:
    if not query:
        return escaped_html
    terms = [
        re.escape(term)
        for term in re.findall(r"[\w\u0600-\u06ff]{3,}", query.lower())
        if len(term) >= 3
    ]
    if not terms:
        return escaped_html
    pattern = re.compile(r"(" + "|".join(sorted(set(terms), key=len, reverse=True)) + r")", re.IGNORECASE)
    return pattern.sub(r"<mark>\1</mark>", escaped_html)


def _build_stats_snapshot(index: Dict, point_count: int = 0) -> Dict:
    review_counts = {}
    source_type_counts = {}
    document_type_counts = {}
    conversion_status_counts = {}
    ingest_ready_books = 0
    ingest_ready_pages = 0
    pending_review_books = 0
    rejected_books = 0
    languages = {}

    for record in index.get("files", []):
        language = str(record.get("language", "unknown") or "unknown")
        languages[language] = languages.get(language, 0) + 1

        review_status = str(record.get("review_status", "approved_auto") or "approved_auto")
        review_counts[review_status] = review_counts.get(review_status, 0) + 1
        source_type = infer_source_type(record)
        source_type_counts[source_type] = source_type_counts.get(source_type, 0) + 1
        document_type = infer_document_type(record)
        document_type_counts[document_type] = document_type_counts.get(document_type, 0) + 1
        conversion_status = infer_conversion_status(record)
        conversion_status_counts[conversion_status] = conversion_status_counts.get(conversion_status, 0) + 1

        pages = int(record.get("total_pages", 0) or 0)
        if bool(record.get("ingest_ready", True)):
            ingest_ready_books += 1
            ingest_ready_pages += pages
        if review_status == "pending_review":
            pending_review_books += 1
        elif review_status == "rejected":
            rejected_books += 1

    return {
        "total_books": int(index.get("total_files", len(index.get("files", []))) or len(index.get("files", []))),
        "ingest_ready_books": ingest_ready_books,
        "pending_review_books": pending_review_books,
        "rejected_books": rejected_books,
        "total_pages": sum(int(r.get("total_pages", 0) or 0) for r in index.get("files", [])),
        "ingest_ready_pages": ingest_ready_pages,
        "total_points_indexed": point_count,
        "languages": languages,
        "review_status_counts": review_counts,
        "source_type_counts": source_type_counts,
        "document_type_counts": document_type_counts,
        "conversion_status_counts": conversion_status_counts,
    }


def _rebuild_index_summary(index: Dict) -> Dict:
    snapshot = _build_stats_snapshot(index)
    index["total_files"] = snapshot["total_books"]
    index["languages"] = dict(sorted(snapshot["languages"].items()))
    index["source_types"] = dict(sorted(snapshot["source_type_counts"].items()))
    index["document_types"] = dict(sorted(snapshot["document_type_counts"].items()))
    index["conversion_status_counts"] = dict(sorted(snapshot["conversion_status_counts"].items()))
    return index


def _refresh_lexical_cache() -> None:
    try:
        if os.path.exists(LEXICAL_INDEX_PATH):
            os.remove(LEXICAL_INDEX_PATH)
    except OSError:
        pass
    try:
        retriever_module._lexical_index = None
    except Exception:
        pass


def _book_review_payload(action: str, reviewed_by: str, note: str) -> Dict:
    now = datetime.now(timezone.utc).isoformat()
    payload = {
        "review_status": "pending_review",
        "review_required": True,
        "review_route": "manual_or_lease_coordinator",
        "reviewed_by": "",
        "reviewed_at": "",
        "review_note": note,
        "ingest_ready": False,
        "page_review_status": "",
        "page_reviewed_by": "",
        "page_reviewed_at": "",
        "page_review_note": "",
    }

    if action == "approved_manual":
        payload.update({
            "review_status": "approved_manual",
            "review_required": False,
            "review_route": reviewed_by or "manual",
            "reviewed_by": reviewed_by or "manual",
            "reviewed_at": now,
            "ingest_ready": True,
        })
    elif action == "approved_lease":
        payload.update({
            "review_status": "approved_lease",
            "review_required": False,
            "review_route": reviewed_by or "lease_coordinator",
            "reviewed_by": reviewed_by or "lease_coordinator",
            "reviewed_at": now,
            "ingest_ready": True,
        })
    elif action == "rejected":
        payload.update({
            "review_status": "rejected",
            "review_required": False,
            "review_route": reviewed_by or "manual",
            "reviewed_by": reviewed_by or "manual",
            "reviewed_at": now,
            "ingest_ready": False,
        })
    elif action in {"pending_review", "duplicate"}:
        payload.update({
            "review_status": "pending_review" if action == "pending_review" else "rejected",
            "review_required": action == "pending_review",
            "review_route": "manual_or_lease_coordinator",
            "reviewed_by": "",
            "reviewed_at": "",
            "ingest_ready": False,
        })
        if action == "duplicate" and not payload["review_note"]:
            payload["review_note"] = note or "duplicate"
    return payload


def _apply_review_to_book_json(book: Dict, record: Dict, action: str, reviewed_by: str, note: str) -> Dict:
    updates = _book_review_payload(action, reviewed_by, note)
    updated = dict(book)
    updated.update({k: v for k, v in updates.items() if not k.startswith("page_")})
    if "review_note" not in updated or not updated["review_note"]:
        updated["review_note"] = note
    return updated


def _update_page_review(book: Dict, page_num: int, action: str, reviewed_by: str, note: str, promote_book: bool) -> Dict:
    pages = list(book.get("pages", []))
    target = None
    for page in pages:
        if int(page.get("page", -1)) == int(page_num):
            target = page
            break
    if target is None:
        raise HTTPException(status_code=404, detail="page not found")

    now = datetime.now(timezone.utc).isoformat()
    page_status = "reviewed" if action in {"page_reviewed", "reviewed"} else "pending_review"
    target["page_review_status"] = page_status
    target["page_reviewed_by"] = reviewed_by
    target["page_reviewed_at"] = now if page_status == "reviewed" else ""
    target["page_review_note"] = note

    page_reviews = dict(book.get("page_reviews", {}))
    page_reviews[str(page_num)] = {
        "page": int(page_num),
        "status": page_status,
        "reviewed_by": reviewed_by,
        "reviewed_at": now if page_status == "reviewed" else "",
        "note": note,
    }
    book["page_reviews"] = page_reviews

    if promote_book:
        book["review_status"] = "approved_manual"
        book["review_required"] = False
        book["review_route"] = reviewed_by or "manual"
        book["reviewed_by"] = reviewed_by or "manual"
        book["reviewed_at"] = now
        book["review_note"] = note or "page reviewed and promoted"
        book["ingest_ready"] = True
    return book


def _delete_book_everywhere(book_id: str) -> Dict:
    loaded = _load_book_record(book_id)
    if not loaded:
        raise HTTPException(status_code=404, detail="not found")

    record = loaded["record"]
    book = loaded["book"]
    json_path = loaded["json_path"]
    index = _load_index()
    content_index = _load_json_file(_content_index_path(), {"entries": [], "total_files": 0})

    try:
        client = QdrantClient(path=QDRANT_PATH)
        state = load_state()
        release_book(state, client, book_id)
        save_state(state)
    except Exception:
        pass

    try:
        if os.path.exists(json_path):
            os.remove(json_path)
    except OSError:
        pass

    index["files"] = [
        item for item in index.get("files", [])
        if (item.get("book_id") or os.path.splitext(item.get("filename", ""))[0]) != book_id
        and os.path.normpath(resolve_index_json_path(item, JSON_DIR)) != os.path.normpath(json_path)
    ]
    _rebuild_index_summary(index)
    _save_json_file(_index_path(), index)

    if isinstance(content_index, dict):
        entries = [
            entry for entry in content_index.get("entries", [])
            if entry.get("book_id") != book_id and os.path.normpath(str(entry.get("json_path", ""))) != os.path.normpath(record.get("json_path", ""))
        ]
        content_index["entries"] = entries
        content_index["total_files"] = len(entries)
        _save_json_file(_content_index_path(), content_index)

    _refresh_lexical_cache()
    return {
        "deleted": True,
        "book_id": book_id,
        "json_path": json_path,
    }


def _replace_content_index_entry(content_index: Dict, new_entry: Dict) -> Dict:
    entries = content_index.get("entries", [])
    if not isinstance(entries, list):
        entries = []
    new_entries = []
    replaced = False
    target_book_id = new_entry.get("book_id")
    target_json_path = new_entry.get("json_path")
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        entry_book_id = entry.get("book_id")
        entry_json_path = entry.get("json_path")
        if entry_book_id == target_book_id or entry_json_path == target_json_path:
            if not replaced:
                new_entries.append(new_entry)
                replaced = True
            continue
        new_entries.append(entry)
    if not replaced:
        new_entries.append(new_entry)
    content_index["entries"] = new_entries
    content_index["total_files"] = len(new_entries)
    return content_index


def _build_content_index_entry(book: Dict, record: Dict, json_path: str) -> Dict:
    signature = _build_book_signature(book)
    merged_record = _apply_book_to_record(record, book, json_path)
    return {
        "book_id": merged_record.get("book_id") or os.path.splitext(merged_record.get("filename", ""))[0],
        "filename": merged_record.get("filename", ""),
        "json_path": merged_record.get("json_path", os.path.relpath(json_path, JSON_DIR).replace("\\", "/")),
        "title": merged_record.get("title", book.get("title", "")),
        "language": merged_record.get("language", book.get("language", "unknown")),
        "source_path": merged_record.get("source_path", ""),
        "source_relpath": merged_record.get("source_relpath", ""),
        "source_ext": merged_record.get("source_ext", infer_source_ext(merged_record)),
        "source_type": merged_record.get("source_type", infer_source_type(merged_record)),
        "document_type": merged_record.get("document_type", infer_document_type(merged_record)),
        "conversion_status": merged_record.get("conversion_status", infer_conversion_status(merged_record)),
        "quality_status": merged_record.get("quality_status", "ok"),
        "quality_reasons": merged_record.get("quality_reasons", []),
        "quality_warnings": merged_record.get("quality_warnings", []),
        "review_status": merged_record.get("review_status", "approved_auto"),
        "review_required": bool(merged_record.get("review_required", False)),
        "review_route": merged_record.get("review_route", "auto"),
        "reviewed_by": merged_record.get("reviewed_by", ""),
        "reviewed_at": merged_record.get("reviewed_at", ""),
        "review_note": merged_record.get("review_note", ""),
        "ingest_ready": bool(merged_record.get("ingest_ready", True)),
        "page_count": signature["page_count"],
        "page_hashes": signature["page_hashes"],
        "content_hash": signature["content_hash"],
        "text_hash": signature["text_hash"],
        "text_simhash": signature["text_simhash"],
        "size_bytes": int(merged_record.get("size_bytes", signature["size_bytes"]) or signature["size_bytes"]),
        "source_hash": merged_record.get("source_hash", ""),
        "total_pages": int(merged_record.get("total_pages", signature["page_count"]) or signature["page_count"]),
    }


def _resolve_source_path(record: Dict) -> str:
    candidates = []
    source_path = str(record.get("source_path", "") or "").strip()
    if source_path:
        candidates.append(source_path)

    source_root = str(record.get("source_root", "") or "").strip()
    source_relpath = str(record.get("source_relpath", "") or "").strip()
    if source_root and source_relpath:
        candidates.append(os.path.join(source_root, source_relpath))
    if source_relpath:
        candidates.append(source_relpath)
    filename = str(record.get("filename", "") or "").strip()
    if filename:
        candidates.append(filename)

    for candidate in candidates:
        if candidate and os.path.exists(candidate):
            return os.path.abspath(candidate)
    return ""


def _source_split_pages_from_text(text: str) -> List[str]:
    if not text or not text.strip():
        return []
    pages = [page.strip() for page in text.split("\f")]
    return [page for page in pages if page]


def _source_html_to_pages(html_text: str) -> List[str]:
    soup = BeautifulSoup(html_text, "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text("\n")
    pages = _source_split_pages_from_text(text)
    return pages or [text.strip()]


def _source_read_text_file(path: Path) -> List[str]:
    with path.open("r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    pages = _source_split_pages_from_text(content)
    return pages or [content.strip()]


def _source_read_html_file(path: Path) -> List[str]:
    return _source_html_to_pages(path.read_text(encoding="utf-8", errors="replace"))


def _source_read_docx_file(path: Path) -> List[str]:
    doc = Document(str(path))
    pages: List[str] = []
    current: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            current.append(text)
        para_xml = para._p.xml
        if ("w:type=\"page\"" in para_xml or "w:lastRenderedPageBreak" in para_xml) and current:
            pages.append("\n".join(current).strip())
            current = []

    if current:
        pages.append("\n".join(current).strip())

    if not pages:
        all_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        return _source_split_pages_from_text(all_text) or [all_text.strip()]

    return pages


def _source_read_pdf_file(path: Path) -> List[str]:
    text = pdf_extract_text(str(path))
    pages = _source_split_pages_from_text(text)
    return pages or [text.strip()]


def _source_read_zip_xhtml_file(path: Path) -> List[str]:
    pages: List[str] = []
    with zipfile.ZipFile(path) as zf:
        members = sorted(
            name for name in zf.namelist()
            if name.lower().endswith((".xhtml", ".html", ".htm"))
        )
        for name in members:
            with zf.open(name) as handle:
                raw = handle.read().decode("utf-8", errors="replace")
            pages.extend(_source_html_to_pages(raw))
    return [page for page in pages if page.strip()]


def _source_read_doc_with_libreoffice(path: Path) -> List[str]:
    if shutil.which("libreoffice") is None:
        raise RuntimeError("libreoffice not available for .doc fallback")

    with tempfile.TemporaryDirectory(prefix="rag-source-lo-") as tmp:
        tmpdir = Path(tmp)
        result = subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "html",
                "--outdir",
                str(tmpdir),
                str(path),
            ],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            check=False,
        )
        html_files = sorted(tmpdir.glob("*.html"))
        if result.returncode != 0 or not html_files:
            raise RuntimeError(
                f"LibreOffice conversion failed for {path.name}: {result.stderr.strip() or result.stdout.strip()}"
            )
        html_text = html_files[0].read_text(encoding="utf-8", errors="replace")
        pages = _source_html_to_pages(html_text)
        return pages or [html_text.strip()]


@lru_cache(maxsize=256)
def _load_source_pages_cached(source_path: str, source_type: str, source_mtime_ns: int, source_size: int) -> Tuple[str, ...]:
    path = Path(source_path)
    if not path.exists():
        return tuple()

    source_type = (source_type or "").strip().lower()
    try:
        if source_type == "txt":
            pages = _source_read_text_file(path)
        elif source_type == "html":
            pages = _source_read_html_file(path)
        elif source_type == "docx":
            pages = _source_read_docx_file(path)
        elif source_type == "pdf":
            pages = _source_read_pdf_file(path)
        elif source_type in {"epub", "ibooks"}:
            pages = _source_read_zip_xhtml_file(path)
        elif source_type == "doc":
            pages = _source_read_doc_with_libreoffice(path)
        else:
            suffix = path.suffix.lower()
            if suffix == ".txt":
                pages = _source_read_text_file(path)
            elif suffix in {".htm", ".html"}:
                pages = _source_read_html_file(path)
            elif suffix == ".docx":
                pages = _source_read_docx_file(path)
            elif suffix == ".pdf":
                pages = _source_read_pdf_file(path)
            elif suffix in {".epub", ".ibooks"}:
                pages = _source_read_zip_xhtml_file(path)
            elif suffix == ".doc":
                pages = _source_read_doc_with_libreoffice(path)
            else:
                pages = [path.read_text(encoding="utf-8", errors="replace").strip()]
    except Exception:
        return tuple()

    return tuple(page for page in pages if page and page.strip())


def _load_source_pages(record: Dict) -> Tuple[List[str], str, str]:
    source_path = _resolve_source_path(record)
    if not source_path:
        return [], infer_source_type(record), ""

    try:
        stat = os.stat(source_path)
        source_mtime_ns = int(stat.st_mtime_ns)
        source_size = int(stat.st_size)
    except OSError:
        return [], infer_source_type(record), source_path

    source_type = infer_source_type(record)
    pages = list(_load_source_pages_cached(source_path, source_type, source_mtime_ns, source_size))
    return pages, source_type, source_path


def _source_page_text(record: Dict, page_num: int) -> Tuple[str, int, int, str, str]:
    pages, source_type, source_path = _load_source_pages(record)
    total_pages = len(pages)
    if total_pages <= 0:
        return "", max(1, int(page_num) or 1), 0, source_type, source_path
    page_index = min(max(int(page_num) - 1, 0), total_pages - 1)
    return pages[page_index], page_index + 1, total_pages, source_type, source_path


def _build_source_nav_html(book_id: str, current_page: int, total_pages: int, theme: str, font_size: int, query: str, target: str = "/sources") -> str:
    if total_pages <= 0:
        return ""
    nav_pages = _compact_page_nav({"total_pages": total_pages}, current_page)
    items = []
    last_num = None
    query_param = quote(query)
    for pnum in nav_pages:
        if last_num is not None and pnum != last_num + 1:
            items.append("<span class='toc-gap'>…</span>")
        active = " active" if pnum == current_page else ""
        items.append(
            f"<a class='toc-item{active}' href='{target}/{quote(str(book_id))}/pages/{pnum}?theme={theme}&font={font_size}&q={query_param}'>{pnum}</a>"
        )
        last_num = pnum
    return "".join(items)


def _render_source_preview_html(record: Dict, book: Dict, page_num: int, theme: str = "dark", font_size: int = 18, query: str = "") -> str:
    book_id = record.get("book_id") or os.path.splitext(record.get("filename", ""))[0]
    title = book.get("title", record.get("title", "Tanpa judul"))
    source_page_text, source_page_num, source_total_pages, source_type, source_path = _source_page_text(record, page_num)
    theme = "light" if str(theme).lower() == "light" else "dark"
    font_size = max(14, min(int(font_size or 18), 28))
    body_bg = "#f7f5ef" if theme == "light" else "linear-gradient(180deg, #0b0f14 0%, #11161d 100%)"
    card_bg = "linear-gradient(180deg, rgba(255,255,255,.97), rgba(247,244,236,.97))" if theme == "light" else "linear-gradient(180deg, rgba(22,27,34,.98), rgba(18,23,31,.98))"
    text_color = "#1d2430" if theme == "light" else "#e6edf3"
    muted_color = "#64748b" if theme == "light" else "#9fb0c3"
    line_color = "rgba(15,23,42,.12)" if theme == "light" else "rgba(255,255,255,.08)"
    panel2 = "#eef2f7" if theme == "light" else "#1f2630"
    content_color = "#111827" if theme == "light" else "#eef3f7"
    prev_page = source_page_num - 1 if source_page_num > 1 else None
    next_page = source_page_num + 1 if source_page_num < source_total_pages else None
    query_param = quote(query)
    page_content_html = html_lib.escape(source_page_text or "").replace("\n", "<br>")
    page_content_html = _highlight_terms_html(page_content_html, query)
    if not page_content_html.strip():
        page_content_html = "<span class='empty'>Tidak ada teks sumber pada halaman ini.</span>"
    prev_link = f"<a class='secondary' href='/sources/{quote(str(book_id))}/pages/{prev_page}?theme={theme}&font={font_size}&q={query_param}'>Halaman sebelumnya</a>" if prev_page else ""
    next_link = f"<a class='secondary' href='/sources/{quote(str(book_id))}/pages/{next_page}?theme={theme}&font={font_size}&q={query_param}'>Halaman berikutnya</a>" if next_page else ""
    json_link = f"/books/{quote(str(book_id))}/pages/{source_page_num}/view?theme={theme}&font={font_size}&q={query_param}"
    review_link = f"/books/{quote(str(book_id))}/pages/{source_page_num}/review?theme={theme}&font={font_size}&q={query_param}"
    nav_html = _build_source_nav_html(book_id, source_page_num, source_total_pages, theme, font_size, query, target="/sources")
    jump_html = f"""
      <form class="jump" onsubmit="const p=this.page.value; if(!p) return false; window.location='/sources/{quote(str(book_id))}/pages/'+encodeURIComponent(p)+'?theme={theme}&font={font_size}&q={query_param}'; return false;">
        <label for="jump-page">Lompat halaman</label>
        <input id="jump-page" name="page" type="number" min="1" max="{source_total_pages}" value="{source_page_num}">
        <button type="submit">Buka</button>
      </form>
    """
    source_kind = "dokumen" if infer_document_type(record) == "html_document" else "buku"
    source_note = "HTML article / dokumen web" if infer_document_type(record) == "html_document" else "Sumber asli"
    if source_total_pages <= 0:
        source_note = "Pratinjau sumber tidak tersedia"
    elif source_total_pages <= 1 and infer_document_type(record) == "html_document":
        source_note = "Dokumen HTML tanpa paginasi stabil"
    page_counter = f"{source_page_num} / {source_total_pages}" if source_total_pages > 0 else f"{source_page_num}"
    return """<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title} - Sumber Halaman {page_num}</title>
<style>
:root {{
  --bg:{body_bg};
  --panel:{card_bg};
  --text:{text_color};
  --muted:{muted_color};
  --line:{line_color};
  --accent:#7dd3fc;
}}
* {{ box-sizing:border-box; }}
html,body {{ margin:0; min-height:100%; }}
body {{
  font-family: ui-serif, Georgia, "Times New Roman", serif;
  background: var(--bg);
  color:var(--text);
}}
.wrap {{ max-width: 1020px; margin: 0 auto; padding: 26px 16px 44px; }}
.card {{
  background: {card_bg};
  border:1px solid var(--line);
  border-radius:24px;
  box-shadow: 0 24px 70px rgba(0,0,0,.35);
  overflow:hidden;
}}
.head {{ padding: 26px 26px 18px; border-bottom:1px solid var(--line); }}
.eyebrow {{
  color: var(--accent);
  letter-spacing: .14em;
  text-transform: uppercase;
  font-size: 12px;
  margin-bottom: 10px;
}}
h1 {{
  margin:0 0 10px;
  font-size: clamp(24px, 4vw, 38px);
  line-height:1.15;
  font-family: ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
}}
.meta {{
  display:flex;
  flex-wrap:wrap;
  gap:8px 12px;
  color:var(--muted);
  font-size: 13px;
  line-height:1.5;
}}
.pill {{
  border:1px solid var(--line);
  border-radius:999px;
  padding:5px 10px;
  background: rgba(255,255,255,.03);
}}
.toolbar {{
  display:flex;
  flex-wrap:wrap;
  gap:10px;
  padding: 16px 26px;
  border-bottom:1px solid var(--line);
  background: rgba(255,255,255,.02);
}}
.toolbar a {{
  text-decoration:none;
  color:#071018;
  background: linear-gradient(135deg, #7dd3fc, #a7f3d0);
  border-radius: 12px;
  padding: 10px 14px;
  font-family: ui-sans-serif, system-ui, sans-serif;
  font-size: 14px;
  font-weight:700;
}}
.toolbar a.secondary {{
  color: var(--text);
  background: var(--panel2);
  border:1px solid var(--line);
}}
.toolbar form.jump {{
  display:flex;
  flex-wrap:wrap;
  gap:8px;
  align-items:center;
  margin-left:auto;
  color: var(--muted);
  font-family: ui-sans-serif, system-ui, sans-serif;
  font-size: 13px;
}}
.toolbar form.jump input {{
  width: 90px;
  border-radius: 10px;
  border:1px solid var(--line);
  background: transparent;
  color: var(--text);
  padding: 9px 10px;
  font-size: 14px;
}}
.toolbar form.jump button {{
  border:none;
  border-radius: 10px;
  padding: 10px 14px;
  background: linear-gradient(135deg, #7dd3fc, #a7f3d0);
  color:#071018;
  font-weight:700;
  cursor:pointer;
}}
.content {{
  padding: 28px 26px 34px;
  font-size: {font_size}px;
  line-height: 1.9;
  white-space: pre-wrap;
  word-break: break-word;
  color: {content_color};
  font-family: ui-serif, Georgia, "Times New Roman", serif;
}}
.content mark {{
  background: rgba(120,215,255,.18);
  color: inherit;
  padding: 0 .14em;
  border-radius: 4px;
}}
.footer {{
  padding: 16px 26px 26px;
  border-top:1px solid var(--line);
  color:var(--muted);
  font-size:13px;
  font-family: ui-sans-serif, system-ui, sans-serif;
}}
.empty {{ color: var(--muted); }}
.toc {{
  display:flex;
  flex-wrap:wrap;
  gap:8px;
  margin-top:10px;
}}
.toc-item {{
  display:inline-flex;
  align-items:center;
  justify-content:center;
  min-width: 34px;
  padding: 7px 9px;
  border-radius: 10px;
  border:1px solid var(--line);
  background: var(--panel2);
  color: var(--text);
  text-decoration:none;
  font-size: 13px;
  font-family: ui-sans-serif, system-ui, sans-serif;
}}
.toc-item.active {{
  background: linear-gradient(135deg, #7dd3fc, #a7f3d0);
  color:#071018;
  font-weight:800;
}}
.toc-gap {{
  color: var(--muted);
  font-family: ui-sans-serif, system-ui, sans-serif;
  font-size: 14px;
  padding: 7px 2px;
}}
@media (max-width: 720px) {{
  .toolbar form.jump {{ margin-left: 0; width: 100%; }}
  .toolbar form.jump input {{ flex: 1 1 110px; }}
}}
</style>
</head>
<body>
<div class="wrap">
  <div class="card">
    <div class="head">
      <div class="eyebrow">Sumber Asli</div>
      <h1>{title}</h1>
      <div class="meta">
        <span class="pill">Book ID: {book_id}</span>
        <span class="pill">Halaman sumber {page_counter}</span>
        <span class="pill">Jenis: {source_kind}</span>
        <span class="pill">{source_note}</span>
        <span class="pill">File: {filename}</span>
      </div>
    </div>
    <div class="toolbar">
      <a href="/books/{book_id}">Detail buku</a>
      <a class="secondary" href="{json_link}">JSON halaman</a>
      <a class="secondary" href="{review_link}">Review workspace</a>
      <a class="secondary" href="/sources/{book_id}/pages/{source_page_num}?theme={theme}&font={font_size}&q={query_param}">Refresh</a>
      <a class="secondary" href="/sources/{book_id}/pages/{source_page_num}?theme={theme_toggle}&font={font_size}&q={query_param}">Tema {theme_toggle}</a>
      <a class="secondary" href="/sources/{book_id}/pages/{source_page_num}?theme={theme}&font={font_sm}&q={query_param}">A-</a>
      <a class="secondary" href="/sources/{book_id}/pages/{source_page_num}?theme={theme}&font={font_md}&q={query_param}">A</a>
      <a class="secondary" href="/sources/{book_id}/pages/{source_page_num}?theme={theme}&font={font_lg}&q={query_param}">A+</a>
      {prev_link}
      {next_link}
      {jump_html}
    </div>
    <div class="content">{page_content_html}</div>
    <div class="footer">
      Sumber asli: {source_path}
      <div class="toc">{nav_html}</div>
    </div>
  </div>
</div>
</body>
</html>""".format(
        title=html_lib.escape(str(title)),
        book_id=html_lib.escape(str(book_id)),
        page_num=page_num,
        source_page_num=source_page_num,
        source_total_pages=source_total_pages,
        page_counter=html_lib.escape(str(page_counter)),
        source_kind=html_lib.escape(str(source_kind)),
        source_note=html_lib.escape(str(source_note)),
        filename=html_lib.escape(str(record.get("filename", "-"))),
        body_bg=body_bg,
        card_bg=card_bg,
        text_color=text_color,
        muted_color=muted_color,
        line_color=line_color,
        panel2=panel2,
        content_color=content_color,
        theme=theme,
        theme_toggle=("light" if theme == "dark" else "dark"),
        font_size=font_size,
        font_sm=max(14, font_size - 1),
        font_md=font_size,
        font_lg=min(28, font_size + 2),
        query_param=query_param,
        prev_link=prev_link,
        next_link=next_link,
        jump_html=jump_html,
        page_content_html=page_content_html,
        source_path=html_lib.escape(str(source_path or "-")),
        json_link=json_link,
        review_link=review_link,
        nav_html=nav_html,
    )


def _render_page_review_html(record: Dict, book: Dict, page: Dict, page_num: int, theme: str = "dark", font_size: int = 18, q: str = "") -> str:
    book_id = record.get("book_id") or os.path.splitext(record.get("filename", ""))[0]
    title = book.get("title", record.get("title", "Tanpa judul"))
    page_content = page.get("content", "")
    source_page_text, source_page_num, source_total_pages, source_type, source_path = _source_page_text(record, page_num)
    theme = "light" if str(theme).lower() == "light" else "dark"
    font_size = max(14, min(int(font_size or 18), 28))
    preview_url = f"/sources/{quote(str(book_id))}/pages/{source_page_num}?theme={theme}&font={font_size}&q={quote(q)}"
    current_page = int(page.get("page", page_num) or page_num)
    prev_page = current_page - 1 if current_page > 1 else None
    next_page = current_page + 1 if current_page < int(book.get("total_pages", current_page) or current_page) else None
    page_review_status = str(page.get("page_review_status", "") or "")
    page_reviewed_by = str(page.get("page_reviewed_by", "") or "")
    page_reviewed_at = str(page.get("page_reviewed_at", "") or "")
    book_review_status = str(book.get("review_status", "approved_auto") or "approved_auto")
    review_note = str(page.get("page_review_note", "") or "")
    current_path = html_lib.escape(str(_resolve_source_path(record) or record.get("source_path", "") or ""))
    json_path = html_lib.escape(str(record.get("json_path", "") or ""))
    return f"""<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Review Halaman - {html_lib.escape(str(title))}</title>
<style>
:root {{
  --bg:#0f1115;
  --panel:#161a22;
  --panel-2:#1d2330;
  --text:#eef2ff;
  --muted:#9aa4b2;
  --line:rgba(255,255,255,.08);
  --accent:#78d7ff;
  --accent-2:#a6ffcb;
  --shadow:0 20px 60px rgba(0,0,0,.35);
}}
*{{box-sizing:border-box}}
html,body{{margin:0;min-height:100%}}
body{{
  font-family:Inter,ui-sans-serif,system-ui,-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;
  color:var(--text);
  background:
    radial-gradient(circle at top left, rgba(120,215,255,.18), transparent 32%),
    radial-gradient(circle at top right, rgba(166,255,203,.12), transparent 28%),
    linear-gradient(180deg,#0b0d11 0%, #11151c 45%, #0b0d11 100%);
}}
.wrap{{max-width:1480px;margin:0 auto;padding:24px 18px 48px}}
.hero{{
  display:grid;
  gap:16px;
  padding:26px;
  border:1px solid var(--line);
  border-radius:24px;
  background:linear-gradient(180deg, rgba(22,26,34,.95), rgba(17,21,28,.9));
  box-shadow:var(--shadow);
}}
.eyebrow{{font-size:12px;letter-spacing:.16em;text-transform:uppercase;color:var(--accent)}}
h1{{margin:0;font-size:clamp(28px,5vw,44px);line-height:1.1}}
.lead{{margin:0;max-width:90ch;color:var(--muted);line-height:1.6;font-size:14px}}
.stats{{display:flex;flex-wrap:wrap;gap:10px}}
.pill{{border:1px solid var(--line);border-radius:999px;padding:8px 12px;background:rgba(255,255,255,.03);font-size:13px}}
.toolbar{{display:flex;flex-wrap:wrap;gap:10px;align-items:center}}
.toolbar a,.toolbar button{{
  border-radius:14px;
  border:none;
  background:linear-gradient(135deg,var(--accent),var(--accent-2));
  color:#071018;
  font-weight:700;
  padding:12px 16px;
  cursor:pointer;
  text-decoration:none;
}}
.toolbar a.secondary{{background:var(--panel-2);color:var(--text);border:1px solid var(--line)}}
.toolbar input,.toolbar select{{
  border-radius:14px;
  border:1px solid var(--line);
  background:rgba(255,255,255,.03);
  color:var(--text);
  padding:12px 14px;
  font-size:14px;
}}
.workspace{{
  display:grid;
  grid-template-columns:minmax(0, 1.18fr) minmax(420px, .82fr);
  gap:16px;
  margin-top:18px;
  align-items:start;
}}
.panel{{
  border:1px solid var(--line);
  border-radius:20px;
  background:linear-gradient(180deg, rgba(255,255,255,.04), rgba(255,255,255,.02));
  box-shadow:0 8px 30px rgba(0,0,0,.18);
  overflow:hidden;
}}
.panel-head{{padding:16px 18px;border-bottom:1px solid var(--line);display:flex;justify-content:space-between;gap:10px;align-items:flex-start;flex-wrap:wrap}}
.panel-head h2{{margin:0;font-size:16px}}
.panel-head .small{{color:var(--muted);font-size:13px;line-height:1.5}}
.preview-iframe{{
  width:100%;
  height: 76vh;
  border:0;
  background:#fff;
}}
.editor-body{{padding:18px;display:grid;gap:12px}}
textarea, input[type="text"]{{
  width:100%;
  border-radius:16px;
  border:1px solid var(--line);
  background:rgba(0,0,0,.25);
  color:var(--text);
  padding:14px 16px;
  font-size:14px;
  outline:none;
}}
textarea{{min-height:42vh;resize:vertical;line-height:1.65;font-family:ui-monospace,SFMono-Regular,Menlo,Monaco,Consolas,"Liberation Mono","Courier New",monospace;}}
textarea:focus,input[type="text"]:focus{{border-color:rgba(120,215,255,.4);box-shadow:0 0 0 4px rgba(120,215,255,.12)}}
.group{{display:grid;gap:6px}}
.label{{color:var(--muted);font-size:12px;letter-spacing:.08em;text-transform:uppercase}}
.buttons{{display:flex;flex-wrap:wrap;gap:8px}}
.buttons button,.buttons a{{
  border-radius:12px;
  border:none;
  background:linear-gradient(135deg,var(--accent),var(--accent-2));
  color:#071018;
  font-weight:700;
  padding:10px 14px;
  cursor:pointer;
  text-decoration:none;
  font-size:13px;
}}
.buttons button.secondary,.buttons a.secondary{{background:var(--panel-2);color:var(--text);border:1px solid var(--line)}}
.buttons button.danger,.buttons a.danger{{background:linear-gradient(135deg,#fda4af,#fecaca)}}
.status{{color:var(--muted);font-size:13px;line-height:1.5}}
.draft{{min-height:14vh}}
.readout{{white-space:pre-wrap;font-family:ui-monospace,SFMono-Regular,Menlo,Monaco,Consolas,"Liberation Mono","Courier New",monospace;font-size:13px;line-height:1.6;padding:14px;border-radius:16px;border:1px solid var(--line);background:rgba(0,0,0,.18);color:#dfe9f7}}
@media (max-width: 1140px){{
  .workspace{{grid-template-columns:1fr}}
  .preview-iframe{{height:58vh}}
}}
</style>
</head>
<body>
<div class="wrap">
  <section class="hero">
    <div class="eyebrow">Review Halaman</div>
    <h1>{html_lib.escape(str(title))}</h1>
    <p class="lead">Panel kiri menampilkan sumber asli halaman, panel kanan menampilkan teks JSON yang bisa diedit langsung. Simpan edit akan menandai buku kembali sebagai pending review agar tidak langsung diingest.</p>
    <div class="stats">
      <div class="pill">Book ID: {html_lib.escape(str(book_id))}</div>
      <div class="pill">Halaman: {current_page}</div>
      <div class="pill">Review buku: {html_lib.escape(book_review_status)}</div>
      <div class="pill">Review halaman: {html_lib.escape(page_review_status or '-')}</div>
      <div class="pill">Sumber: {html_lib.escape(str(source_type))}</div>
    </div>
    <div class="toolbar">
      <a href="/library">Daftar Buku</a>
      <a class="secondary" href="/books/{quote(str(book_id))}/pages/{current_page}/view?theme={theme}&font={font_size}&q={quote(q)}">JSON halaman</a>
      <a class="secondary" href="/sources/{quote(str(book_id))}/pages/{source_page_num}?theme={theme}&font={font_size}&q={quote(q)}">Buka sumber</a>
      <a class="secondary" href="/books/{quote(str(book_id))}/edit">Edit JSON buku</a>
      <a class="secondary" href="/books/{quote(str(book_id))}/raw" target="_blank" rel="noreferrer">Raw JSON</a>
      <input id="pageJump" type="text" value="{current_page}" aria-label="Lompat halaman">
      <button id="goPage" type="button">Buka halaman</button>
    </div>
  </section>

  <div class="workspace">
    <section class="panel">
      <div class="panel-head">
        <div>
          <h2>Preview sumber asli</h2>
          <div class="small">Path: {current_path}<br>JSON: {json_path}</div>
        </div>
        <div class="buttons">
          <a class="secondary" href="{preview_url}" target="_blank" rel="noreferrer">Buka di tab baru</a>
          <a class="secondary" href="/sources/{quote(str(book_id))}/pages/{source_page_num}?theme={theme}&font={font_size}&q={quote(q)}">Refresh sumber</a>
        </div>
      </div>
      <iframe class="preview-iframe" src="{preview_url}" title="Preview sumber asli"></iframe>
    </section>

    <section class="panel">
      <div class="panel-head">
        <div>
          <h2>Editor halaman</h2>
          <div class="small">Edit isi halaman ini, lalu simpan. Setelah itu status akan kembali pending review.</div>
        </div>
        <div class="buttons">
          <a class="secondary" href="/books/{quote(str(book_id))}/pages/{current_page}/view?theme={theme}&font={font_size}&q={quote(q)}">Baca JSON</a>
          <a class="secondary" href="/books/{quote(str(book_id))}/pages/{current_page}/review?theme={theme}&font={font_size}&q={quote(q)}">Refresh</a>
        </div>
      </div>
      <div class="editor-body">
        <div class="group">
          <div class="label">Isi halaman JSON</div>
          <textarea id="pageContent">{html_lib.escape(page_content)}</textarea>
        </div>
        <div class="group">
          <div class="label">Catatan review</div>
          <input id="reviewNote" type="text" value="{html_lib.escape(review_note)}" placeholder="Catatan singkat, misal: noise footer sudah dibuang">
        </div>
        <div class="group">
          <div class="label">Permintaan repair draft</div>
          <input id="repairInstruction" type="text" placeholder="Contoh: perbaiki Arabic yang rusak tanpa mengubah makna">
        </div>
        <div class="group">
          <div class="label">Marker replacement</div>
          <input id="markerHint" type="text" value="[[FIX_QS 5:41]] / [[FIX_HADITH bukhari:1]] / [[DELETE_START]]...[[DELETE_END]] / [[DORAR_SEARCH ...]]" readonly>
        </div>
        <div class="buttons">
          <button id="savePage" type="button">Simpan halaman</button>
          <button id="markReviewed" type="button">Halaman reviewed</button>
          <button id="markPending" type="button" class="secondary">Halaman pending</button>
          <button id="approveBook" type="button">Buku siap ingest</button>
          <button id="duplicateBook" type="button" class="secondary">Tolak dobel</button>
          <button id="deleteBook" type="button" class="danger">Hapus buku</button>
        </div>
        <div class="buttons">
          <button id="repairLocal" type="button" class="secondary">Draft repair lokal</button>
          <button id="repairLease" type="button" class="secondary">Draft repair lease</button>
          <button id="applyDraft" type="button" class="secondary">Pakai draft</button>
        </div>
        <div class="buttons">
          <button id="applyMarkers" type="button" class="secondary">Preview markers</button>
          <button id="applyMarkerDraft" type="button" class="secondary">Pakai hasil marker</button>
          <label style="display:inline-flex;align-items:center;gap:8px;color:var(--muted);font-size:13px;">
            <input id="autoDorarFirst" type="checkbox">
            Auto Dorar pertama
          </label>
        </div>
        <div class="group">
          <div class="label">Cari hadits lokal</div>
          <div style="display:flex;gap:8px;flex-wrap:wrap;">
            <input id="localHadithQuery" type="text" placeholder="Contoh: إنما الأعمال بالنيات" style="flex:1 1 260px;">
            <input id="localHadithCollection" type="text" placeholder="Koleksi opsional: bukhari, muslim, nawawi40" style="flex:1 1 220px;">
            <button id="searchLocalHadith" type="button" class="secondary">Cari Lokal</button>
          </div>
          <div class="status">Pencarian lokal memakai dataset hadits yang sudah diimpor. Kosongkan koleksi untuk cari semua koleksi.</div>
          <div id="localHadithResults" class="readout" style="min-height:12vh;"></div>
        </div>
        <div class="group">
          <div class="label">Cari hadits di Dorar</div>
          <div style="display:flex;gap:8px;flex-wrap:wrap;">
            <input id="dorarQuery" type="text" placeholder="Contoh: إنما الأعمال بالنيات" style="flex:1 1 280px;">
            <button id="searchDorar" type="button" class="secondary">Cari Dorar</button>
          </div>
          <div class="status">Klik kandidat untuk menyisipkan teks ke posisi cursor di editor halaman.</div>
          <div id="dorarResults" class="readout" style="min-height:12vh;"></div>
        </div>
        <div class="group">
          <div class="label">Preview hasil marker</div>
          <textarea id="markerOutput" class="draft" placeholder="Hasil preview marker replacement akan muncul di sini" readonly></textarea>
          <div id="markerDiff" class="readout" style="min-height:12vh;">Belum ada preview marker.</div>
        </div>
        <div id="status" class="status">Siap review.</div>
        <div class="group">
          <div class="label">Draft perbaikan</div>
          <textarea id="repairOutput" class="draft" placeholder="Hasil draft perbaikan akan muncul di sini"></textarea>
        </div>
        <div class="group">
          <div class="label">Review metadata</div>
          <div class="readout">page_review_status: {html_lib.escape(page_review_status or '-')}\npage_reviewed_by: {html_lib.escape(page_reviewed_by or '-')}\npage_reviewed_at: {html_lib.escape(page_reviewed_at or '-')}\nbook_review_status: {html_lib.escape(book_review_status)}\nsource_path: {current_path}\nsource_type: {html_lib.escape(str(source_type))}</div>
        </div>
      </div>
    </section>
  </div>
</div>
<script>
(function() {{
  const api = window.location.origin;
  const bookId = {json.dumps(str(book_id))};
  const pageNum = {int(current_page)};
  const currentTheme = {json.dumps(theme)};
  const currentFont = {int(font_size)};
  const currentQuery = {json.dumps(q)};
  const pageContent = document.getElementById('pageContent');
  const reviewNote = document.getElementById('reviewNote');
  const repairInstruction = document.getElementById('repairInstruction');
  const repairOutput = document.getElementById('repairOutput');
  const markerOutput = document.getElementById('markerOutput');
  const markerDiff = document.getElementById('markerDiff');
  const markerHint = document.getElementById('markerHint');
  const dorarQuery = document.getElementById('dorarQuery');
  const dorarResults = document.getElementById('dorarResults');
  const localHadithQuery = document.getElementById('localHadithQuery');
  const localHadithCollection = document.getElementById('localHadithCollection');
  const localHadithResults = document.getElementById('localHadithResults');
  const autoDorarFirst = document.getElementById('autoDorarFirst');
  const status = document.getElementById('status');
  const pageJump = document.getElementById('pageJump');
  function setStatus(text) {{
    status.textContent = text;
  }}
  function insertAtCursor(textarea, text) {{
    const start = textarea.selectionStart ?? textarea.value.length;
    const end = textarea.selectionEnd ?? textarea.value.length;
    const before = textarea.value.slice(0, start);
    const after = textarea.value.slice(end);
    textarea.value = `${{before}}${{text}}${{after}}`;
    const pos = start + text.length;
    textarea.focus();
    textarea.selectionStart = pos;
    textarea.selectionEnd = pos;
  }}
  function escapeInline(text) {{
    return String(text || '')
      .replaceAll('&', '&amp;')
      .replaceAll('<', '&lt;')
      .replaceAll('>', '&gt;')
      .replaceAll('"', '&quot;')
      .replaceAll("'", '&#39;');
  }}
  function renderDorarCandidates(map) {{
    const entries = Object.entries(map || {{}});
    if (!entries.length) {{
      dorarResults.innerHTML = '<div class="empty">Belum ada hasil Dorar.</div>';
      return;
    }}
    dorarResults.innerHTML = entries.map(([query, items]) => {{
      const selectedIndex = Number.isInteger(currentDorarChoices[query]) ? currentDorarChoices[query] : -1;
      const rows = (items || []).map((item, idx) => {{
        const text = escapeInline(item.text || '');
        const source = escapeInline(item.source || '');
        const grade = escapeInline(item.grade || '');
        const isSelected = idx === selectedIndex;
        return `
          <div style="border-top:1px solid var(--line);padding-top:10px;margin-top:10px;">
            <div style="font-size:12px;color:var(--muted);margin-bottom:6px;">#${{idx + 1}} ${{source ? `· ${{source}}` : ''}} ${{grade ? `· ${{grade}}` : ''}}</div>
            <div style="white-space:pre-wrap;line-height:1.6;">${{text || '(kosong)'}} </div>
            <div class="buttons" style="margin-top:8px;">
              <button type="button" class="secondary" data-insert-dorar="${{encodeURIComponent(query)}}" data-index="${{idx}}">Sisipkan</button>
              <button type="button" class="${{isSelected ? 'primary' : 'secondary'}}" data-pick-dorar="${{encodeURIComponent(query)}}" data-index="${{idx}}">${{isSelected ? 'Terpilih' : 'Pilih kandidat'}}</button>
            </div>
          </div>
        `;
      }}).join('');
      return `
        <div style="margin-bottom:14px;">
          <div style="font-weight:700;margin-bottom:6px;">${{escapeInline(query)}}</div>
          <div style="color:var(--muted);font-size:13px;">${{items.length}} kandidat${{selectedIndex >= 0 ? ` · terpilih #${{selectedIndex + 1}}` : ''}}</div>
          ${{rows}}
        </div>
      `;
    }}).join('');
    dorarResults.querySelectorAll('[data-insert-dorar]').forEach(btn => {{
      btn.addEventListener('click', () => {{
        const query = decodeURIComponent(btn.dataset.insertDorar || '');
        const idx = parseInt(btn.dataset.index || '0', 10);
        const entry = Object.entries(currentDorarMap).find(([k]) => k === query);
        const items = entry ? entry[1] : [];
        const item = items[idx];
        if (!item) return;
        insertAtCursor(pageContent, item.text || '');
        setStatus(`Kandidat Dorar disisipkan untuk query: ${{query}}`);
      }});
    }});
    dorarResults.querySelectorAll('[data-pick-dorar]').forEach(btn => {{
      btn.addEventListener('click', async () => {{
        const query = decodeURIComponent(btn.dataset.pickDorar || '');
        const idx = parseInt(btn.dataset.index || '0', 10);
        currentDorarChoices[query] = idx;
        renderDorarCandidates(currentDorarMap);
        setStatus(`Kandidat Dorar dipilih untuk query: ${{query}} (#${{idx + 1}}).`);
        try {{
          await previewMarkers();
        }} catch (err) {{
          setStatus(`Preview marker gagal: ${{err.message}}`);
        }}
      }});
    }});
  }}
  let currentDorarMap = {{}};
  let currentDorarChoices = {{}};
  let currentLocalHadithMap = {{}};
  async function postJson(url, body) {{
    const resp = await fetch(url, {{
      method: 'POST',
      headers: {{ 'Content-Type': 'application/json' }},
      body: JSON.stringify(body)
    }});
    const data = await resp.json().catch(() => ({{}}));
    if (!resp.ok) {{
      throw new Error(data.detail || data.error || `HTTP ${{resp.status}}`);
    }}
    return data;
  }}
  async function savePage() {{
    setStatus('Menyimpan halaman...');
    const data = await postJson(`${{api}}/books/${{encodeURIComponent(bookId)}}/pages/${{pageNum}}/edit`, {{
      content: pageContent.value,
      reviewed_by: 'web',
      note: reviewNote.value || ''
    }});
    setStatus(`Tersimpan. Review status: ${{data.review_status || 'pending_review'}}`);
  }}
  async function sendReview(action, scope, opts={{}}) {{
    const data = await postJson(`${{api}}/books/${{encodeURIComponent(bookId)}}/review`, {{
      scope,
      action,
      reviewed_by: opts.reviewed_by || 'web',
      note: opts.note || '',
      page_num: scope === 'page' ? pageNum : null,
      promote_book: !!opts.promote_book,
      delete_physical: !!opts.delete_physical,
    }});
    return data;
  }}
  async function runRepair(backend) {{
    const note = repairInstruction.value || '';
    setStatus('Membuat draft repair...');
    const data = await postJson(`${{api}}/books/${{encodeURIComponent(bookId)}}/pages/${{pageNum}}/repair`, {{
      backend,
      instruction: note,
      note,
      reviewed_by: 'web'
    }});
    repairOutput.value = data.suggested_content || '';
    setStatus(`Draft siap dari backend ${{data.backend_used || backend}}.`);
  }}
  function renderMarkerPreview(data) {{
    const resolved = data.resolved_content || '';
    markerOutput.value = resolved;
    const diffText = data.diff_text || 'Tidak ada diff.';
    markerDiff.innerHTML = `
      <div style="margin-bottom:8px;">${{escapeInline(data.diff_summary || 'Preview siap.')}}</div>
      <pre style="white-space:pre-wrap;line-height:1.5;margin:0;">${{escapeInline(diffText)}}</pre>
    `;
  }}
  function renderLocalHadithCandidates(map) {{
    const entries = Object.entries(map || {{}});
    if (!entries.length) {{
      localHadithResults.innerHTML = '<div class="empty">Belum ada hasil hadits lokal.</div>';
      return;
    }}
    localHadithResults.innerHTML = entries.map(([query, items]) => {{
      const rows = (items || []).map((item, idx) => {{
        const text = escapeInline(item.text || '');
        const source = escapeInline(item.source || '');
        const grade = escapeInline(item.grade || '');
        return `
          <div style="border-top:1px solid var(--line);padding-top:10px;margin-top:10px;">
            <div style="font-size:12px;color:var(--muted);margin-bottom:6px;">#${{idx + 1}} ${{source ? `· ${{source}}` : ''}} ${{grade ? `· ${{grade}}` : ''}}</div>
            <div style="white-space:pre-wrap;line-height:1.6;">${{text || '(kosong)'}} </div>
            <div class="buttons" style="margin-top:8px;">
              <button type="button" class="secondary" data-insert-local-hadith="${{encodeURIComponent(query)}}" data-index="${{idx}}">Sisipkan</button>
            </div>
          </div>
        `;
      }}).join('');
      return `
        <div style="margin-bottom:14px;">
          <div style="font-weight:700;margin-bottom:6px;">${{escapeInline(query)}}</div>
          <div style="color:var(--muted);font-size:13px;">${{items.length}} hasil lokal</div>
          ${{rows}}
        </div>
      `;
    }}).join('');
    localHadithResults.querySelectorAll('[data-insert-local-hadith]').forEach(btn => {{
      btn.addEventListener('click', () => {{
        const query = decodeURIComponent(btn.dataset.insertLocalHadith || '');
        const idx = parseInt(btn.dataset.index || '0', 10);
        const entry = Object.entries(currentLocalHadithMap).find(([k]) => k === query);
        const items = entry ? entry[1] : [];
        const item = items[idx];
        if (!item) return;
        insertAtCursor(pageContent, item.text || '');
        setStatus(`Kandidat hadits lokal disisipkan untuk query: ${{query}}`);
      }});
    }});
  }}
  async function previewMarkers() {{
    setStatus('Membuat preview marker...');
    const data = await postJson(`${{api}}/books/${{encodeURIComponent(bookId)}}/pages/${{pageNum}}/apply-markers`, {{
      content: pageContent.value,
      dorar_policy: autoDorarFirst.checked ? 'first' : 'preserve',
      dorar_choices: currentDorarChoices
    }});
    currentDorarMap = data.dorar_candidates || {{}};
    renderDorarCandidates(currentDorarMap);
    renderMarkerPreview(data);
    const unresolved = (data.unresolved || []).length;
    setStatus(`Preview marker selesai. stages: ${{(data.stages || []).length}}, unresolved: ${{unresolved}}`);
    return data;
  }}
  async function applyMarkers() {{
    return await previewMarkers();
  }}
  function applyMarkerDraftToEditor() {{
    if (!markerOutput.value.trim()) {{
      setStatus('Preview marker masih kosong.');
      return;
    }}
    if (!confirm('Pakai hasil preview marker ke editor halaman?')) return;
    pageContent.value = markerOutput.value;
    setStatus('Hasil preview marker sudah dipakai ke editor.');
  }}
  async function searchDorar() {{
    const query = (dorarQuery.value || '').trim();
    if (!query) {{
      setStatus('Masukkan query Dorar dulu.');
      return;
    }}
    setStatus('Mencari Dorar...');
    const data = await postJson(`${{api}}/admin/hadith/dorar/search`, {{
      query,
      limit: 5
    }});
    currentDorarMap = {{ [query]: data.results || [] }};
    renderDorarCandidates(currentDorarMap);
    setStatus(`Dorar menemukan ${{(data.results || []).length}} kandidat untuk query: ${{query}}`);
  }}
  async function searchLocalHadith() {{
    const query = (localHadithQuery.value || '').trim();
    const collection = (localHadithCollection.value || '').trim();
    if (!query) {{
      setStatus('Masukkan query hadits lokal dulu.');
      return;
    }}
    setStatus('Mencari hadits lokal...');
    const data = await postJson(`${{api}}/admin/hadith/local/search`, {{
      query,
      collection,
      limit: 10
    }});
    currentLocalHadithMap = {{ [query]: data.results || [] }};
    renderLocalHadithCandidates(currentLocalHadithMap);
    setStatus(`Hadits lokal menemukan ${{(data.results || []).length}} kandidat untuk query: ${{query}}`);
  }}
  document.getElementById('savePage').addEventListener('click', async () => {{
    try {{
      await savePage();
      window.location.reload();
    }} catch (err) {{
      setStatus(`Gagal menyimpan: ${{err.message}}`);
    }}
  }});
  document.getElementById('markReviewed').addEventListener('click', async () => {{
    try {{
      setStatus('Menandai halaman reviewed...');
      await sendReview('page_reviewed', 'page', {{ reviewed_by: 'web', note: reviewNote.value || '', promote_book: false }});
      window.location.reload();
    }} catch (err) {{
      setStatus(`Gagal: ${{err.message}}`);
    }}
  }});
  document.getElementById('markPending').addEventListener('click', async () => {{
    try {{
      setStatus('Menandai halaman pending...');
      await sendReview('page_pending', 'page', {{ reviewed_by: 'web', note: reviewNote.value || '', promote_book: false }});
      window.location.reload();
    }} catch (err) {{
      setStatus(`Gagal: ${{err.message}}`);
    }}
  }});
  document.getElementById('approveBook').addEventListener('click', async () => {{
    try {{
      setStatus('Menandai buku siap ingest...');
      await sendReview('approved_manual', 'book', {{ reviewed_by: 'web', note: reviewNote.value || '', promote_book: true }});
      window.location.reload();
    }} catch (err) {{
      setStatus(`Gagal: ${{err.message}}`);
    }}
  }});
  document.getElementById('duplicateBook').addEventListener('click', async () => {{
    try {{
      setStatus('Menandai duplikat...');
      await sendReview('duplicate', 'book', {{ reviewed_by: 'web', note: reviewNote.value || 'duplicate', delete_physical: false }});
      window.location.reload();
    }} catch (err) {{
      setStatus(`Gagal: ${{err.message}}`);
    }}
  }});
  document.getElementById('deleteBook').addEventListener('click', async () => {{
    if (!confirm('Hapus fisik JSON + hilangkan dari index/Qdrant?')) return;
    try {{
      setStatus('Menghapus buku...');
      await sendReview('delete', 'book', {{ reviewed_by: 'web', note: reviewNote.value || '', delete_physical: true }});
      window.location = '/library';
    }} catch (err) {{
      setStatus(`Gagal: ${{err.message}}`);
    }}
  }});
  document.getElementById('repairLocal').addEventListener('click', async () => {{
    try {{
      await runRepair('local');
    }} catch (err) {{
      setStatus(`Repair gagal: ${{err.message}}`);
    }}
  }});
  document.getElementById('repairLease').addEventListener('click', async () => {{
    try {{
      await runRepair('lease');
    }} catch (err) {{
      setStatus(`Repair gagal: ${{err.message}}`);
    }}
  }});
  document.getElementById('applyDraft').addEventListener('click', async () => {{
    if (!repairOutput.value.trim()) {{
      setStatus('Draft masih kosong.');
      return;
    }}
    if (!confirm('Ganti isi editor dengan draft perbaikan?')) return;
    pageContent.value = repairOutput.value;
    setStatus('Draft diterapkan ke editor.');
  }});
  document.getElementById('applyMarkers').addEventListener('click', async () => {{
    try {{
      await applyMarkers();
    }} catch (err) {{
      setStatus(`Apply marker gagal: ${{err.message}}`);
    }}
  }});
  document.getElementById('applyMarkerDraft').addEventListener('click', async () => {{
    try {{
      applyMarkerDraftToEditor();
    }} catch (err) {{
      setStatus(`Pakai hasil marker gagal: ${{err.message}}`);
    }}
  }});
  document.getElementById('searchDorar').addEventListener('click', async () => {{
    try {{
      await searchDorar();
    }} catch (err) {{
      setStatus(`Cari Dorar gagal: ${{err.message}}`);
    }}
  }});
  document.getElementById('searchLocalHadith').addEventListener('click', async () => {{
    try {{
      await searchLocalHadith();
    }} catch (err) {{
      setStatus(`Cari hadits lokal gagal: ${{err.message}}`);
    }}
  }});
  document.getElementById('goPage').addEventListener('click', () => {{
    const next = parseInt(pageJump.value, 10);
    if (!Number.isFinite(next) || next < 1) return;
    window.location = `/books/${{encodeURIComponent(bookId)}}/pages/${{next}}/review?theme=${{encodeURIComponent(currentTheme)}}&font=${{encodeURIComponent(currentFont)}}&q=${{encodeURIComponent(currentQuery)}}`;
  }});
  pageJump.addEventListener('keydown', (e) => {{
    if (e.key === 'Enter') {{
      document.getElementById('goPage').click();
    }}
  }});
  currentDorarMap = {{}};
  currentDorarChoices = {{}};
  renderDorarCandidates(currentDorarMap);
  currentLocalHadithMap = {{}};
  renderLocalHadithCandidates(currentLocalHadithMap);
  markerOutput.value = '';
  markerDiff.textContent = 'Belum ada preview marker.';
}})();
</script>
</body>
</html>"""


def _update_book_after_json_edit(
    book_id: str,
    record: Dict,
    book: Dict,
    json_path: str,
    *,
    release_indexed: bool = False,
    extra_record_updates: Dict | None = None,
) -> Dict:
    index = _load_index()
    idx_record = _find_index_record(index, book_id)
    if not idx_record:
        raise HTTPException(status_code=404, detail="index record not found")

    updated_record = _apply_book_to_record(idx_record, book, json_path)
    if extra_record_updates:
        updated_record.update(extra_record_updates)
    for i, existing in enumerate(index.get("files", [])):
        existing_book_id = existing.get("book_id") or os.path.splitext(existing.get("filename", ""))[0]
        if existing_book_id == book_id:
            index["files"][i] = updated_record
            break

    _rebuild_index_summary(index)
    _save_json_file(_index_path(), index)

    content_index = _load_json_file(_content_index_path(), {"entries": [], "total_files": 0})
    if isinstance(content_index, dict):
        content_entry = _build_content_index_entry(book, updated_record, json_path)
        content_index = _replace_content_index_entry(content_index, content_entry)
        _save_json_file(_content_index_path(), content_index)

    if release_indexed:
        try:
            client = QdrantClient(path=QDRANT_PATH)
            state = load_state()
            release_book(state, client, book_id)
            save_state(state)
        except Exception:
            pass

    _refresh_lexical_cache()
    return updated_record


def _render_page_html(record: Dict, book: Dict, page: Dict, page_num: int, theme: str = "dark", font_size: int = 19, query: str = "") -> str:
    title = book.get("title", "Tanpa judul")
    book_id = record.get("book_id") or os.path.splitext(record.get("filename", ""))[0]
    page_start = page.get("page", page_num)
    prev_page = page_start - 1 if page_start > 1 else None
    next_page = page_start + 1 if page_start < int(book.get("total_pages", page_start) or page_start) else None
    content = page.get("content", "")
    theme = "light" if str(theme).lower() == "light" else "dark"
    font_size = max(14, min(int(font_size or 19), 28))
    body_bg = "#f7f5ef" if theme == "light" else "linear-gradient(180deg, #0b0f14 0%, #11161d 100%)"
    card_bg = "linear-gradient(180deg, rgba(255,255,255,.96), rgba(247,244,236,.96))" if theme == "light" else "linear-gradient(180deg, rgba(22,27,34,.96), rgba(18,23,31,.96))"
    text_color = "#1d2430" if theme == "light" else "#e6edf3"
    muted_color = "#64748b" if theme == "light" else "#9fb0c3"
    line_color = "rgba(15,23,42,.12)" if theme == "light" else "rgba(255,255,255,.08)"
    panel2 = "#eef2f7" if theme == "light" else "#1f2630"
    content_color = "#111827" if theme == "light" else "#eef3f7"
    prev_link_html = f"<a class='secondary' href='/books/{quote(str(book_id))}/pages/{prev_page}/view?theme={theme}&font={font_size}&q={quote(query)}'>Halaman sebelumnya</a>" if prev_page else ""
    next_link_html = f"<a class='secondary' href='/books/{quote(str(book_id))}/pages/{next_page}/view?theme={theme}&font={font_size}&q={quote(query)}'>Halaman berikutnya</a>" if next_page else ""
    content_html = html_lib.escape(content).replace("\n", "<br>") if content else "<span class='empty'>Tidak ada teks pada halaman ini.</span>"
    content_html = _highlight_terms_html(content_html, query)
    page_review_status = str(page.get("page_review_status", "") or "")
    page_reviewed_by = str(page.get("page_reviewed_by", "") or "")
    page_reviewed_at = str(page.get("page_reviewed_at", "") or "")
    book_review_status = str(book.get("review_status", "approved_auto") or "approved_auto")
    theme_toggle = "light" if theme == "dark" else "dark"
    font_sm = max(14, font_size - 1)
    font_md = font_size
    font_lg = min(28, font_size + 2)
    font_mobile = max(14, font_size - 1)
    max_pages = int(book.get("total_pages", page_start) or page_start)
    query_param = quote(query)
    source_url = f"/sources/{quote(str(book_id))}/pages/{page_start}?theme={theme}&font={font_size}&q={query_param}"
    review_url = f"/books/{quote(str(book_id))}/pages/{page_start}/review?theme={theme}&font={font_size}&q={query_param}"
    page_nav_items = []
    current_nav_pages = _compact_page_nav(book, page_start)
    last_num = None
    for pnum in current_nav_pages:
        if last_num is not None and pnum != last_num + 1:
            page_nav_items.append("<span class='toc-gap'>…</span>")
        active = " active" if pnum == page_start else ""
        page_nav_items.append(
            f"<a class='toc-item{active}' href='/books/{quote(str(book_id))}/pages/{pnum}/view?theme={theme}&font={font_size}&q={query_param}'>{pnum}</a>"
        )
        last_num = pnum
    toc_html = "".join(page_nav_items)
    jump_html = f"""
      <form class="jump" onsubmit="const p=this.page.value; if(!p) return false; window.location='/books/{quote(str(book_id))}/pages/'+encodeURIComponent(p)+'/view?theme={theme}&font={font_size}&q={query_param}'; return false;">
        <label for="jump-page">Lompat halaman</label>
        <input id="jump-page" name="page" type="number" min="1" max="{max_pages}" value="{page_start}">
        <button type="submit">Buka</button>
      </form>
    """
    review_html = f"""
      <div class="reviewbox">
        <div class="reviewline">
          <span class="reviewlabel">Review halaman</span>
          <span class="reviewpill">{html_lib.escape(page_review_status or '-')}</span>
          <span class="reviewmeta">{html_lib.escape(page_reviewed_by or '-')} {html_lib.escape(page_reviewed_at or '')}</span>
        </div>
        <div class="reviewactions">
          <button type="button" data-review-scope="page" data-review-action="page_reviewed">Halaman reviewed</button>
          <button type="button" data-review-scope="page" data-review-action="page_pending">Halaman pending</button>
          <button type="button" data-review-scope="book" data-review-action="approved_manual">Buku siap ingest</button>
          <button type="button" data-review-scope="book" data-review-action="duplicate">Tolak dobel</button>
          <button type="button" class="danger" data-review-scope="book" data-review-action="delete">Hapus buku</button>
          <a class="secondary" href="{review_url}">Review workspace</a>
        </div>
      </div>
    """
    book_id_js = json.dumps(str(book_id))
    prev_url_js = json.dumps(f"/books/{book_id}/pages/{prev_page}/view?theme={theme}&font={font_size}&q={query_param}") if prev_page else "null"
    next_url_js = json.dumps(f"/books/{book_id}/pages/{next_page}/view?theme={theme}&font={font_size}&q={query_param}") if next_page else "null"
    return """<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title} - Halaman {page_start}</title>
<style>
:root {{
  --bg:{body_bg};
  --panel:{card_bg};
  --panel2:{panel2};
  --text:{text_color};
  --muted:{muted_color};
  --accent:#7dd3fc;
  --line:{line_color};
}}
* {{ box-sizing:border-box; }}
html,body {{ margin:0; min-height:100%; }}
body {{
  font-family: ui-serif, Georgia, "Times New Roman", serif;
  background: var(--bg);
  color:var(--text);
}}
.wrap {{ max-width: 980px; margin: 0 auto; padding: 28px 16px 48px; }}
.card {{
  background: {card_bg};
  border:1px solid var(--line);
  border-radius:24px;
  box-shadow: 0 24px 70px rgba(0,0,0,.35);
  overflow:hidden;
}}
.reader {{
  display:grid;
  grid-template-columns: minmax(0, 1fr) 260px;
  gap: 0;
}}
.head {{ padding: 26px 26px 18px; border-bottom:1px solid var(--line); }}
.eyebrow {{
  color: var(--accent);
  letter-spacing: .14em;
  text-transform: uppercase;
  font-size: 12px;
  margin-bottom: 10px;
}}
h1 {{
  margin:0 0 10px;
  font-size: clamp(24px, 4vw, 40px);
  line-height:1.15;
  font-family: ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
}}
.meta {{
  display:flex;
  flex-wrap:wrap;
  gap:8px 12px;
  color:var(--muted);
  font-size: 13px;
  line-height:1.5;
}}
.pill {{
  border:1px solid var(--line);
  border-radius:999px;
  padding:5px 10px;
  background: rgba(255,255,255,.03);
}}
.toolbar {{
  display:flex;
  flex-wrap:wrap;
  gap:10px;
  padding: 16px 26px;
  border-bottom:1px solid var(--line);
  background: rgba(255,255,255,.02);
}}
.toolbar a {{
  text-decoration:none;
  color:#071018;
  background: linear-gradient(135deg, #7dd3fc, #a7f3d0);
  border-radius: 12px;
  padding: 10px 14px;
  font-family: ui-sans-serif, system-ui, sans-serif;
  font-size: 14px;
  font-weight:700;
}}
.toolbar a.secondary {{
  color: var(--text);
  background: var(--panel2);
  border:1px solid var(--line);
}}
.toolbar form.jump {{
  display:flex;
  flex-wrap:wrap;
  gap:8px;
  align-items:center;
  margin-left:auto;
  color: var(--muted);
  font-family: ui-sans-serif, system-ui, sans-serif;
  font-size: 13px;
}}
.toolbar form.jump input {{
  width: 90px;
  border-radius: 10px;
  border:1px solid var(--line);
  background: transparent;
  color: var(--text);
  padding: 9px 10px;
  font-size: 14px;
}}
.toolbar form.jump button {{
  border:none;
  border-radius: 10px;
  padding: 10px 14px;
  background: linear-gradient(135deg, #7dd3fc, #a7f3d0);
  color:#071018;
  font-weight:700;
  cursor:pointer;
}}
.reviewbox {{
  display:grid;
  gap:10px;
  padding: 16px 26px 18px;
  border-bottom:1px solid var(--line);
  background: rgba(255,255,255,.02);
}}
.reviewline {{
  display:flex;
  flex-wrap:wrap;
  gap:8px;
  align-items:center;
  color: var(--muted);
  font-family: ui-sans-serif, system-ui, sans-serif;
  font-size: 13px;
}}
.reviewlabel {{
  color: var(--text);
  font-weight:700;
}}
.reviewpill {{
  display:inline-flex;
  align-items:center;
  border-radius:999px;
  padding:5px 10px;
  border:1px solid var(--line);
  background: var(--panel2);
  color: var(--text);
}}
.reviewactions {{
  display:flex;
  flex-wrap:wrap;
  gap:8px;
}}
.reviewactions button,
.reviewactions a {{
  border:none;
  border-radius: 12px;
  padding: 10px 14px;
  background: linear-gradient(135deg, #7dd3fc, #a7f3d0);
  color:#071018;
  font-weight:700;
  cursor:pointer;
  font-family: ui-sans-serif, system-ui, sans-serif;
  font-size: 13px;
  text-decoration:none;
  display:inline-flex;
  align-items:center;
  justify-content:center;
}}
.reviewactions button.danger {{
  background: linear-gradient(135deg, #fda4af, #fecaca);
}}
.reviewactions a.secondary {{
  background: var(--panel2);
  color: var(--text);
  border:1px solid var(--line);
}}
.sidebar {{
  border-left:1px solid var(--line);
  background: rgba(255,255,255,.02);
  padding: 20px 16px;
}}
.sidebar h2 {{
  margin: 0 0 12px;
  font-size: 15px;
  font-family: ui-sans-serif, system-ui, sans-serif;
}}
.sidebar p {{
  margin: 0 0 12px;
  color: var(--muted);
  font-size: 13px;
  line-height: 1.5;
  font-family: ui-sans-serif, system-ui, sans-serif;
}}
.toc {{
  display:flex;
  flex-wrap:wrap;
  gap:8px;
}}
.toc-item {{
  display:inline-flex;
  align-items:center;
  justify-content:center;
  min-width: 34px;
  padding: 7px 9px;
  border-radius: 10px;
  border:1px solid var(--line);
  background: var(--panel2);
  color: var(--text);
  text-decoration:none;
  font-size: 13px;
  font-family: ui-sans-serif, system-ui, sans-serif;
}}
.toc-item.active {{
  background: linear-gradient(135deg, #7dd3fc, #a7f3d0);
  color:#071018;
  font-weight:800;
}}
.toc-gap {{
  color: var(--muted);
  font-family: ui-sans-serif, system-ui, sans-serif;
  font-size: 14px;
  padding: 7px 2px;
}}
.content {{
  padding: 28px 26px 34px;
  font-size: {font_size}px;
  line-height: 1.9;
  white-space: pre-wrap;
  word-break: break-word;
  color: {content_color};
  font-family: ui-serif, Georgia, "Times New Roman", serif;
}}
.footer {{
  padding: 16px 26px 26px;
  border-top:1px solid var(--line);
  color:var(--muted);
  font-size:13px;
  font-family: ui-sans-serif, system-ui, sans-serif;
}}
.empty {{ color: var(--muted); }}
@media (max-width: 720px) {{
  .reader {{ grid-template-columns: 1fr; }}
  .sidebar {{ border-left:none; border-top:1px solid var(--line); }}
  .content {{ font-size: {font_mobile}px; line-height: 1.8; }}
  .toolbar form.jump {{ margin-left: 0; width: 100%; }}
  .toolbar form.jump input {{ flex: 1 1 110px; }}
}}
</style>
</head>
<body>
<div class="wrap">
  <div class="card">
    <div class="head">
      <div class="eyebrow">Halaman Sumber</div>
      <h1>{title}</h1>
      <div class="meta">
        <span class="pill">Book ID: {book_id}</span>
        <span class="pill">Halaman {page_start}</span>
        <span class="pill">Bahasa: {language}</span>
        <span class="pill">File: {filename}</span>
      </div>
    </div>
    <div class="reader">
      <div>
        <div class="toolbar">
          <a href="/books/{book_id}">Detail buku</a>
          <a class="secondary" href="{source_url}">Buka halaman sumber</a>
          <a class="secondary" href="{review_url}">Review workspace</a>
          <a class="secondary" href="/books/{book_id}/pages/{page_start}">JSON mentah</a>
          <a class="secondary" href="/books/{book_id}/pages/{page_start}/view?theme={theme_toggle}&font={font_size}&q={query_param}">Tema {theme_toggle}</a>
          <a class="secondary" href="/books/{book_id}/pages/{page_start}/view?theme={theme}&font={font_sm}&q={query_param}">A-</a>
          <a class="secondary" href="/books/{book_id}/pages/{page_start}/view?theme={theme}&font={font_md}&q={query_param}">A</a>
          <a class="secondary" href="/books/{book_id}/pages/{page_start}/view?theme={theme}&font={font_lg}&q={query_param}">A+</a>
          {prev_link_html}
          {next_link_html}
          {jump_html}
        </div>
        {review_html}
        <div class="content">{content_html}</div>
        <div class="footer">
          Sumber JSON: {json_path}<br>
          Jalur asli: {source_path}
          <br>Status review buku: {book_review_status}
        </div>
      </div>
      <aside class="sidebar">
        <h2>Daftar Halaman</h2>
        <p>Halaman kecil di sekitar posisi baca saat ini.</p>
        <div class="toc">{toc_html}</div>
      </aside>
    </div>
  </div>
</div>
<script>
(function() {{
  const prevUrl = {prev_url_js};
  const nextUrl = {next_url_js};
  const bookId = {book_id_js};
  const pageNum = {page_start};
  const apiBase = window.location.origin;
  async function sendReview(action, scope, opts = {{}}) {{
    const payload = {{
      scope,
      action,
      reviewed_by: opts.reviewed_by || 'web',
      note: opts.note || '',
      page_num: scope === 'page' ? pageNum : null,
      promote_book: !!opts.promote_book,
      delete_physical: !!opts.delete_physical,
    }};
    const resp = await fetch(`${{apiBase}}/books/${{encodeURIComponent(bookId)}}/review`, {{
      method: 'POST',
      headers: {{ 'Content-Type': 'application/json' }},
      body: JSON.stringify(payload)
    }});
    const data = await resp.json().catch(() => ({{}}));
    if (!resp.ok) {{
      throw new Error(data.detail || data.error || `HTTP ${{resp.status}}`);
    }}
    return data;
  }}
  document.querySelectorAll('[data-review-action]').forEach(btn => {{
    btn.addEventListener('click', async () => {{
      const action = btn.dataset.reviewAction;
      const scope = btn.dataset.reviewScope || 'book';
      const reviewer = prompt('Reviewer', 'web') || 'web';
      const note = prompt('Catatan (boleh kosong)', '') || '';
      const promoteBook = action === 'page_reviewed' ? confirm('Promosikan buku jadi siap ingest jika perlu?') : false;
      const deletePhysical = action === 'delete' ? confirm('Hapus fisik JSON + hilangkan dari index/Qdrant?') : false;
      btn.disabled = true;
      try {{
        await sendReview(action, scope, {{
          reviewed_by: reviewer,
          note,
          promote_book: promoteBook,
          delete_physical: deletePhysical,
        }});
        window.location.reload();
      }} catch (err) {{
        alert(`Gagal menyimpan review: ${{err.message}}`);
      }} finally {{
        btn.disabled = false;
      }}
    }});
  }});
  document.addEventListener('keydown', (e) => {{
    if (e.target && ['INPUT','TEXTAREA','SELECT'].includes(e.target.tagName)) return;
    if (e.key === 'ArrowLeft' || e.key === 'j') {{
      if (prevUrl) window.location = prevUrl;
    }}
    if (e.key === 'ArrowRight' || e.key === 'k') {{
      if (nextUrl) window.location = nextUrl;
    }}
  }});
}})();
</script>
</body>
</html>""".format(
        title=html_lib.escape(str(title)),
        book_id=html_lib.escape(str(book_id)),
        page_start=page_start,
        language=html_lib.escape(str(book.get("language", "-"))),
        filename=html_lib.escape(str(record.get("filename", "-"))),
        body_bg=body_bg,
        card_bg=card_bg,
        text_color=text_color,
        muted_color=muted_color,
        line_color=line_color,
        panel2=panel2,
        content_color=content_color,
        theme=theme,
        theme_toggle=theme_toggle,
        font_size=font_size,
        font_sm=font_sm,
        font_md=font_md,
        font_lg=font_lg,
        font_mobile=font_mobile,
        query_param=query_param,
        jump_html=jump_html,
        source_url=source_url,
        review_url=review_url,
        book_id_js=book_id_js,
        prev_url_js=prev_url_js,
        next_url_js=next_url_js,
        prev_link_html=prev_link_html,
        next_link_html=next_link_html,
        content_html=content_html,
        json_path=html_lib.escape(str(record.get("json_path", "-") or "-")),
        source_path=html_lib.escape(str(record.get("source_path", "-") or "-")),
        toc_html=toc_html,
        review_html=review_html,
        book_review_status=html_lib.escape(book_review_status),
    )


def _render_library_html(page: int = 1, limit: int = 24, q: str = "", lang: str = "", status: str = "") -> str:
    index = _load_index()
    from qdrant_client import QdrantClient
    point_count = 0
    try:
        c = QdrantClient(path=QDRANT_PATH)
        info = c.get_collection(COLLECTION_NAME)
        point_count = info.points_count
    except Exception:
        point_count = 0

    stats = _build_stats_snapshot(index, point_count)
    books = sorted(
        index.get("files", []),
        key=lambda r: (
            str(r.get("language", "unknown")),
            str(r.get("title", "")),
            str(r.get("filename", "")),
        ),
    )
    q_norm = (q or "").strip().lower()
    lang_norm = (lang or "").strip().lower()
    status_norm = (status or "").strip().lower()
    filtered_books = []
    for record in books:
        record_lang = str(record.get("language", "unknown") or "unknown").lower()
        source_type = infer_source_type(record)
        document_type = infer_document_type(record)
        conversion_status = infer_conversion_status(record)
        review_status = str(record.get("review_status", "approved_auto") or "approved_auto").lower()
        quality_status = str(record.get("quality_status", "ok") or "ok").lower()
        ingest_ready = bool(record.get("ingest_ready", True))
        haystack = " ".join([
            str(record.get("book_id") or os.path.splitext(record.get("filename", ""))[0]),
            str(record.get("title", "")),
            str(record.get("filename", "")),
            record_lang,
            source_type,
            document_type,
            conversion_status,
            review_status,
            quality_status,
        ]).lower()
        if q_norm and q_norm not in haystack:
            continue
        if lang_norm and lang_norm != record_lang:
            continue
        if status_norm and status_norm not in {conversion_status.lower(), review_status, quality_status}:
            continue
        filtered_books.append(record)

    total_filtered = len(filtered_books)
    total_pages = max(1, (total_filtered + max(1, limit) - 1) // max(1, limit))
    current_page = max(1, min(int(page or 1), total_pages))
    per_page = max(1, min(int(limit or 24), 96))
    if per_page != limit:
        limit = per_page
        total_pages = max(1, (total_filtered + limit - 1) // limit)
        current_page = max(1, min(current_page, total_pages))
    start = (current_page - 1) * limit
    end = start + limit
    page_books = filtered_books[start:end]

    def _library_url(page_no: int) -> str:
        params = {"page": page_no, "limit": limit}
        if q_norm:
            params["q"] = q
        if lang_norm:
            params["lang"] = lang
        if status_norm:
            params["status"] = status
        return f"/library?{urlencode(params)}"

    prev_url = _library_url(current_page - 1) if current_page > 1 else ""
    next_url = _library_url(current_page + 1) if current_page < total_pages else ""
    page_start_label = start + 1 if total_filtered else 0
    page_end_label = min(end, total_filtered)
    cards_html = []
    for record in page_books:
        book_id = record.get("book_id") or os.path.splitext(record.get("filename", ""))[0]
        title = record.get("title", record.get("filename", "Tanpa judul"))
        language = record.get("language", "unknown")
        source_type = infer_source_type(record)
        document_type = infer_document_type(record)
        conversion_status = infer_conversion_status(record)
        review_status = record.get("review_status", "approved_auto")
        ingest_ready = bool(record.get("ingest_ready", True))
        total_pages = int(record.get("total_pages", 0) or 0)
        first_reader = f"/sources/{quote(str(book_id))}/pages/1" if total_pages > 0 else ""
        first_review = f"/books/{quote(str(book_id))}/pages/1/review" if total_pages > 0 else ""
        editor_url = f"/books/{quote(str(book_id))}/edit"
        raw_url = f"/books/{quote(str(book_id))}/raw"
        dataset = " ".join([
            str(book_id),
            str(title),
            str(language),
            str(source_type),
            str(document_type),
            str(conversion_status),
            str(review_status),
            str(record.get("filename", "")),
        ])
        badges = []
        badges.append(f"<span class='badge'>{html_lib.escape(language)}</span>")
        badges.append(f"<span class='badge'>{html_lib.escape(source_type)}</span>")
        badges.append(f"<span class='badge'>{html_lib.escape(document_type)}</span>")
        badges.append(f"<span class='badge status-{html_lib.escape(conversion_status)}'>{html_lib.escape(conversion_status)}</span>")
        badges.append(f"<span class='badge status-{html_lib.escape(review_status)}'>{html_lib.escape(review_status)}</span>")
        if not ingest_ready:
            badges.append("<span class='badge danger'>not ingest ready</span>")
        read_action = f'<a href="{first_reader}">Baca</a>' if first_reader else '<span class="disabled">Belum ada halaman</span>'
        review_action = f'<a class="secondary" href="{first_review}">Review</a>' if first_review else '<span class="disabled">Belum ada halaman</span>'
        cards_html.append(
            f"""
            <article class="book-card" data-book-id="{html_lib.escape(str(book_id))}" data-search="{html_lib.escape(dataset).lower()}">
              <div class="book-main">
                <div class="book-title">{html_lib.escape(str(title))}</div>
                <div class="book-meta">
                  <span>{html_lib.escape(str(book_id))}</span>
                  <span>{total_pages} pages</span>
                  <span>{html_lib.escape(record.get('filename', '-'))}</span>
                </div>
                <div class="book-badges">{''.join(badges)}</div>
                </div>
                <div class="book-actions">
                {read_action}
                {review_action}
                <a class="secondary" href="{editor_url}">Edit JSON</a>
                <a class="secondary" href="{raw_url}" target="_blank" rel="noreferrer">Raw JSON</a>
                <button type="button" data-review-scope="book" data-review-action="approved_manual">Siap ingest</button>
                <button type="button" class="secondary" data-review-scope="book" data-review-action="duplicate">Tolak dobel</button>
                <button type="button" class="danger" data-review-scope="book" data-review-action="delete">Hapus</button>
              </div>
            </article>
            """
        )

    cards = "\n".join(cards_html)
    return f"""<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Daftar Buku</title>
<style>
:root {{
  --bg:#0f1115;
  --panel:#161a22;
  --panel-2:#1d2330;
  --text:#eef2ff;
  --muted:#9aa4b2;
  --line:rgba(255,255,255,.08);
  --accent:#78d7ff;
  --accent-2:#a6ffcb;
  --shadow:0 20px 60px rgba(0,0,0,.35);
}}
*{{box-sizing:border-box}}
html,body{{margin:0;min-height:100%}}
body{{
  font-family:Inter,ui-sans-serif,system-ui,-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;
  color:var(--text);
  background:
    radial-gradient(circle at top left, rgba(120,215,255,.18), transparent 32%),
    radial-gradient(circle at top right, rgba(166,255,203,.12), transparent 28%),
    linear-gradient(180deg,#0b0d11 0%, #11151c 45%, #0b0d11 100%);
}}
.wrap{{max-width:1280px;margin:0 auto;padding:28px 18px 48px}}
.hero{{
  display:grid;
  gap:16px;
  padding:28px;
  border:1px solid var(--line);
  border-radius:24px;
  background:linear-gradient(180deg, rgba(22,26,34,.95), rgba(17,21,28,.9));
  box-shadow:var(--shadow);
}}
.eyebrow{{font-size:12px;letter-spacing:.16em;text-transform:uppercase;color:var(--accent)}}
h1{{margin:0;font-size:clamp(32px,6vw,58px);line-height:1}}
.lead{{margin:0;max-width:80ch;color:var(--muted);line-height:1.6;font-size:15px}}
.stats{{display:flex;flex-wrap:wrap;gap:10px}}
.pill{{border:1px solid var(--line);border-radius:999px;padding:8px 12px;background:rgba(255,255,255,.03);font-size:13px}}
.toolbar{{display:flex;flex-wrap:wrap;gap:10px;align-items:center}}
.toolbar input,.toolbar select{{
  border-radius:14px;
  border:1px solid var(--line);
  background:rgba(255,255,255,.03);
  color:var(--text);
  padding:12px 14px;
  font-size:14px;
  min-width:180px;
}}
.toolbar a,.toolbar button{{
  border-radius:14px;
  border:none;
  background:linear-gradient(135deg,var(--accent),var(--accent-2));
  color:#071018;
  font-weight:700;
  padding:12px 16px;
  cursor:pointer;
  text-decoration:none;
}}
.toolbar a.secondary{{background:var(--panel-2);color:var(--text);border:1px solid var(--line)}}
.pager{{display:flex;flex-wrap:wrap;gap:12px;justify-content:space-between;align-items:center;margin-top:16px}}
.pager-links{{display:flex;flex-wrap:wrap;gap:10px;align-items:center}}
.pager .disabled{{display:inline-flex;align-items:center;border-radius:12px;padding:10px 14px;color:var(--muted);border:1px dashed var(--line);background:rgba(255,255,255,.02)}}
.grid{{display:grid;gap:14px;margin-top:18px}}
.book-card{{
  display:grid;
  grid-template-columns:minmax(0,1fr) auto;
  gap:16px;
  align-items:flex-start;
  border:1px solid var(--line);
  border-radius:20px;
  background:linear-gradient(180deg, rgba(255,255,255,.04), rgba(255,255,255,.02));
  padding:18px;
  box-shadow:0 8px 30px rgba(0,0,0,.18);
}}
.book-title{{font-size:18px;font-weight:800;line-height:1.35}}
.book-meta{{display:flex;flex-wrap:wrap;gap:10px;margin-top:6px;color:var(--muted);font-size:13px}}
.book-badges{{display:flex;flex-wrap:wrap;gap:8px;margin-top:12px}}
.badge{{
  display:inline-flex;
  align-items:center;
  gap:6px;
  padding:6px 10px;
  border-radius:999px;
  border:1px solid var(--line);
  background:rgba(255,255,255,.03);
  font-size:12px;
  color:var(--text);
}}
.badge.danger{{background:rgba(255,90,90,.16);color:#ffd6d6}}
.badge.status-failed{{background:rgba(255,90,90,.12)}}
.badge.status-good{{background:rgba(90,255,160,.12)}}
.badge.status-pending_review{{background:rgba(255,210,90,.12)}}
.book-actions{{display:flex;flex-wrap:wrap;gap:8px;justify-content:flex-end}}
.book-actions a{{
  border-radius:12px;
  border:none;
  background:linear-gradient(135deg,var(--accent),var(--accent-2));
  color:#071018;
  font-weight:700;
  padding:10px 14px;
  text-decoration:none;
}}
.book-actions a.secondary{{background:var(--panel-2);color:var(--text);border:1px solid var(--line)}}
.book-actions button{{
  border-radius:12px;
  border:none;
  background:linear-gradient(135deg,var(--accent),var(--accent-2));
  color:#071018;
  font-weight:700;
  padding:10px 14px;
  cursor:pointer;
}}
.book-actions button.secondary{{background:var(--panel-2);color:var(--text);border:1px solid var(--line)}}
.book-actions button.danger{{background:linear-gradient(135deg,#fda4af,#fecaca)}}
.book-actions .disabled{{display:inline-flex;align-items:center;border-radius:12px;padding:10px 14px;color:var(--muted);border:1px dashed var(--line);background:rgba(255,255,255,.02)}}
.empty{{padding:22px;border:1px dashed var(--line);border-radius:18px;color:var(--muted);background:rgba(255,255,255,.02)}}
@media (max-width: 820px){{
  .book-card{{grid-template-columns:1fr}}
  .book-actions{{justify-content:flex-start}}
  .toolbar input,.toolbar select{{min-width:140px;flex:1 1 140px}}
}}
</style>
</head>
<body>
<div class="wrap">
  <section class="hero">
    <div class="eyebrow">Daftar Buku</div>
    <h1>Koleksi yang tersedia</h1>
    <p class="lead">Pilih buku untuk membaca, membuka editor JSON, atau melihat raw JSON. Daftar ini dipaginasi agar halaman awal cepat dan ringan.</p>
    <div class="stats">
      <div class="pill">Total buku: {stats['total_books']}</div>
      <div class="pill">Siap ingest: {stats['ingest_ready_books']}</div>
      <div class="pill">Pending review: {stats['pending_review_books']}</div>
      <div class="pill">Rejected: {stats['rejected_books']}</div>
      <div class="pill">Points: {stats['total_points_indexed']}</div>
    </div>
    <form class="toolbar" method="get" action="/library">
      <input name="q" type="search" placeholder="Cari judul, book_id, bahasa, status..." value="{html_lib.escape(q)}" />
      <select name="lang">
        <option value="">Semua bahasa</option>
        {''.join(f'<option value="{html_lib.escape(lang)}"{" selected" if str(lang).lower() == lang_norm else ""}>{html_lib.escape(lang)} ({count})</option>' for lang, count in sorted(stats["languages"].items()))}
      </select>
      <select name="status">
        <option value="">Semua status</option>
        <option value="good"{" selected" if status_norm == "good" else ""}>good</option>
        <option value="failed"{" selected" if status_norm == "failed" else ""}>failed</option>
        <option value="pending_review"{" selected" if status_norm == "pending_review" else ""}>pending_review</option>
        <option value="approved_auto"{" selected" if status_norm == "approved_auto" else ""}>approved_auto</option>
        <option value="approved_manual"{" selected" if status_norm == "approved_manual" else ""}>approved_manual</option>
        <option value="approved_lease"{" selected" if status_norm == "approved_lease" else ""}>approved_lease</option>
      </select>
      <select name="limit">
        <option value="12"{" selected" if limit == 12 else ""}>12 / halaman</option>
        <option value="24"{" selected" if limit == 24 else ""}>24 / halaman</option>
        <option value="48"{" selected" if limit == 48 else ""}>48 / halaman</option>
        <option value="96"{" selected" if limit == 96 else ""}>96 / halaman</option>
      </select>
      <button type="submit">Terapkan</button>
      <a href="/library" class="secondary">Reset</a>
      <a href="/" class="secondary">Kembali ke Search</a>
      <a class="secondary" href="/stats" target="_blank" rel="noreferrer">Stats JSON</a>
    </form>
    <div class="pager">
      <div class="pill">Menampilkan {page_start_label}-{page_end_label} dari {total_filtered} buku</div>
      <div class="pager-links">
        {f'<a class="secondary" href="{prev_url}">Sebelumnya</a>' if prev_url else '<span class="disabled">Sebelumnya</span>'}
        <span class="pill">Halaman {current_page}/{total_pages}</span>
        {f'<a class="secondary" href="{next_url}">Berikutnya</a>' if next_url else '<span class="disabled">Berikutnya</span>'}
      </div>
    </div>
  </section>
  <section class="grid" id="bookGrid">
    {cards if cards else '<div class="empty">Tidak ada buku.</div>'}
  </section>
  <div class="pager">
    <div class="pager-links">
      {f'<a class="secondary" href="{prev_url}">Sebelumnya</a>' if prev_url else '<span class="disabled">Sebelumnya</span>'}
      <span class="pill">Halaman {current_page}/{total_pages}</span>
      {f'<a class="secondary" href="{next_url}">Berikutnya</a>' if next_url else '<span class="disabled">Berikutnya</span>'}
    </div>
  </div>
</div>
<script>
(function() {{
  const apiBase = window.location.origin;
  async function sendReview(bookId, action) {{
    const reviewer = prompt('Reviewer', 'web') || 'web';
    const note = prompt('Catatan (boleh kosong)', '') || '';
    const deletePhysical = action === 'delete' ? confirm('Hapus fisik JSON + hilangkan dari index/Qdrant?') : false;
    const resp = await fetch(`${{apiBase}}/books/${{encodeURIComponent(bookId)}}/review`, {{
      method: 'POST',
      headers: {{ 'Content-Type': 'application/json' }},
      body: JSON.stringify({{
        scope: 'book',
        action,
        reviewed_by: reviewer,
        note,
        delete_physical: deletePhysical,
      }})
    }});
    const data = await resp.json().catch(() => ({{}}));
    if (!resp.ok) {{
      throw new Error(data.detail || data.error || `HTTP ${{resp.status}}`);
    }}
    window.location.reload();
  }}
  document.querySelectorAll('[data-review-action]').forEach(btn => {{
    btn.addEventListener('click', async () => {{
      const action = btn.dataset.reviewAction;
      const bookCard = btn.closest('.book-card');
      const bookId = bookCard?.dataset.bookId || '';
      if (!bookId) return;
      btn.disabled = true;
      try {{
        await sendReview(bookId, action);
      }} catch (err) {{
        alert(`Gagal menyimpan review: ${{err.message}}`);
      }} finally {{
        btn.disabled = false;
      }}
    }});
  }});
}})();
</script>
</body>
</html>"""


def _render_editor_html(book_id: str, record: Dict, book: Dict) -> str:
    raw_json = html_lib.escape(json.dumps(book, ensure_ascii=False, indent=2, sort_keys=True))
    title = html_lib.escape(str(book.get("title", record.get("title", "Tanpa judul"))))
    json_path = html_lib.escape(str(record.get("json_path", "")))
    source_path = html_lib.escape(str(record.get("source_path", "")))
    status = html_lib.escape(str(record.get("review_status", "approved_auto")))
    return f"""<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Edit JSON - {title}</title>
<style>
:root {{
  --bg:#0f1115;
  --panel:#161a22;
  --panel-2:#1d2330;
  --text:#eef2ff;
  --muted:#9aa4b2;
  --line:rgba(255,255,255,.08);
  --accent:#78d7ff;
  --accent-2:#a6ffcb;
  --shadow:0 20px 60px rgba(0,0,0,.35);
}}
*{{box-sizing:border-box}}
html,body{{margin:0;min-height:100%}}
body{{
  font-family:Inter,ui-sans-serif,system-ui,-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;
  color:var(--text);
  background:
    radial-gradient(circle at top left, rgba(120,215,255,.18), transparent 32%),
    radial-gradient(circle at top right, rgba(166,255,203,.12), transparent 28%),
    linear-gradient(180deg,#0b0d11 0%, #11151c 45%, #0b0d11 100%);
}}
.wrap{{max-width:1280px;margin:0 auto;padding:28px 18px 48px}}
.hero{{
  display:grid;
  gap:16px;
  padding:28px;
  border:1px solid var(--line);
  border-radius:24px;
  background:linear-gradient(180deg, rgba(22,26,34,.95), rgba(17,21,28,.9));
  box-shadow:var(--shadow);
}}
.eyebrow{{font-size:12px;letter-spacing:.16em;text-transform:uppercase;color:var(--accent)}}
h1{{margin:0;font-size:clamp(28px,5vw,44px);line-height:1.1}}
.lead{{margin:0;max-width:90ch;color:var(--muted);line-height:1.6;font-size:14px}}
.stats{{display:flex;flex-wrap:wrap;gap:10px}}
.pill{{border:1px solid var(--line);border-radius:999px;padding:8px 12px;background:rgba(255,255,255,.03);font-size:13px}}
.toolbar{{display:flex;flex-wrap:wrap;gap:10px;align-items:center}}
.toolbar a,.toolbar button{{
  border-radius:14px;
  border:none;
  background:linear-gradient(135deg,var(--accent),var(--accent-2));
  color:#071018;
  font-weight:700;
  padding:12px 16px;
  cursor:pointer;
  text-decoration:none;
}}
.toolbar a.secondary{{background:var(--panel-2);color:var(--text);border:1px solid var(--line)}}
.editor-wrap{{
  display:grid;
  grid-template-columns:minmax(0,1fr) 320px;
  gap:14px;
  margin-top:18px;
}}
.editor-card,.side-card{{
  border:1px solid var(--line);
  border-radius:20px;
  background:linear-gradient(180deg, rgba(255,255,255,.04), rgba(255,255,255,.02));
  box-shadow:0 8px 30px rgba(0,0,0,.18);
}}
.editor-card{{padding:18px}}
.side-card{{padding:18px;display:grid;gap:10px;align-content:start}}
.editor-card textarea{{
  width:100%;
  min-height:72vh;
  resize:vertical;
  border-radius:16px;
  border:1px solid var(--line);
  background:rgba(0,0,0,.25);
  color:var(--text);
  padding:16px;
  font-family:ui-monospace,SFMono-Regular,Menlo,Monaco,Consolas,"Liberation Mono","Courier New",monospace;
  font-size:13px;
  line-height:1.55;
  outline:none;
}}
.editor-card textarea:focus{{border-color:rgba(120,215,255,.4);box-shadow:0 0 0 4px rgba(120,215,255,.12)}}
.status{{color:var(--muted);font-size:13px;line-height:1.5}}
.warn{{padding:12px 14px;border-radius:14px;border:1px solid rgba(255,200,90,.25);background:rgba(255,200,90,.08);color:#ffe6ae;font-size:13px;line-height:1.5}}
.field{{display:grid;gap:4px}}
.field label{{color:var(--muted);font-size:12px;letter-spacing:.08em;text-transform:uppercase}}
.field .value{{font-size:14px;word-break:break-word}}
@media (max-width: 980px){{
  .editor-wrap{{grid-template-columns:1fr}}
  .editor-card textarea{{min-height:60vh}}
}}
</style>
</head>
<body>
<div class="wrap">
  <section class="hero">
    <div class="eyebrow">JSON Editor</div>
    <h1>{title}</h1>
    <p class="lead">Edit JSON buku secara langsung. Simpan akan memperbarui file JSON, _index.json, dan _content_index.json. Jika Anda mengubah isi halaman, search yang memakai embedding tetap perlu reindex agar sinkron.</p>
    <div class="stats">
      <div class="pill">Book ID: {html_lib.escape(book_id)}</div>
      <div class="pill">Status review: {status}</div>
      <div class="pill">JSON: {json_path}</div>
    </div>
    <div class="toolbar">
      <a href="/library">Kembali ke daftar buku</a>
      <a class="secondary" href="/sources/{quote(str(book_id))}/pages/1">Baca buku</a>
      <a class="secondary" href="/books/{quote(str(book_id))}/raw" target="_blank" rel="noreferrer">Raw JSON</a>
      <button id="saveBtn" type="button">Simpan JSON</button>
    </div>
  </section>
  <div class="editor-wrap">
    <section class="editor-card">
      <textarea id="jsonEditor" spellcheck="false">{raw_json}</textarea>
    </section>
    <aside class="side-card">
      <div class="warn">Perubahan disimpan ke JSON corpus. Jika Anda mengubah isi halaman, jalankan reindex/ingest ulang agar pencarian embedding dan indeks lokal ikut menyesuaikan.</div>
      <div class="field">
        <label>Source path</label>
        <div class="value">{source_path}</div>
      </div>
      <div class="field">
        <label>Reviewer</label>
        <div class="value">{html_lib.escape(str(record.get('reviewed_by', '') or '-'))}</div>
      </div>
      <div class="field">
        <label>Ingest ready</label>
        <div class="value">{'yes' if bool(record.get('ingest_ready', True)) else 'no'}</div>
      </div>
      <div id="status" class="status">Siap edit.</div>
    </aside>
  </div>
</div>
<script>
(function() {{
  const editor = document.getElementById('jsonEditor');
  const status = document.getElementById('status');
  const saveBtn = document.getElementById('saveBtn');
  const api = window.location.origin;
  saveBtn.addEventListener('click', async () => {{
    saveBtn.disabled = true;
    status.textContent = 'Menyimpan...';
    try {{
      const resp = await fetch(`${{api}}/books/{quote(str(book_id))}/edit`, {{
        method: 'POST',
        headers: {{ 'Content-Type': 'application/json' }},
        body: JSON.stringify({{ json_text: editor.value }})
      }});
      const data = await resp.json().catch(() => ({{}}));
      if (!resp.ok) {{
        throw new Error(data.detail || data.error || `HTTP ${{resp.status}}`);
      }}
      status.textContent = `Tersimpan: ${{data.json_path || 'ok'}}`;
    }} catch (err) {{
      status.textContent = `Gagal menyimpan: ${{err.message}}`;
    }} finally {{
      saveBtn.disabled = false;
    }}
  }});
}})();
</script>
</body>
</html>"""


class AskRequest(BaseModel):
    query: str
    top_k: int = 5
    language: str = "id"
    strict: bool = True
    mode: str = "auto"


class RetrieveRequest(BaseModel):
    query: str
    top_k: int = 5
    language: str = "id"
    mode: str = "balanced"


class JsonEditRequest(BaseModel):
    json_text: str


class PageEditRequest(BaseModel):
    content: str
    reviewed_by: str = "web"
    note: str = ""


class PageRepairRequest(BaseModel):
    backend: str = "local"
    instruction: str = ""
    note: str = ""
    reviewed_by: str = "web"


class PageMarkerApplyRequest(BaseModel):
    content: str
    dorar_policy: str = "preserve"
    dorar_choices: Dict[str, int] = Field(default_factory=dict)


class DorarSearchRequest(BaseModel):
    query: str
    limit: int = 5


class LocalHadithSearchRequest(BaseModel):
    query: str
    limit: int = 10
    collection: str = ""


class ReviewActionRequest(BaseModel):
    scope: str = "book"
    action: str
    reviewed_by: str = "web"
    note: str = ""
    page_num: int | None = None
    promote_book: bool = False
    delete_physical: bool = False


class AskResponse(BaseModel):
    answer: str
    backend_used: str
    mode: str
    sources: List[Dict]


class SourceInfo(BaseModel):
    title: str
    page_start: int
    page_end: int
    book_id: str
    filename: str


@app.get("/health")
def health():
    storage_path = os.path.join(QDRANT_PATH, "collection", COLLECTION_NAME, "storage.sqlite")
    state_path = os.path.join(QDRANT_PATH, "collection", COLLECTION_NAME)
    qdrant_status = "unknown"
    if os.path.exists(storage_path):
        try:
            size_mb = os.path.getsize(storage_path) / (1024 * 1024)
            qdrant_status = f"present ({size_mb:.1f} MB)"
        except Exception:
            qdrant_status = "present"
    elif os.path.exists(state_path):
        qdrant_status = "present"

    return {
        "status": "ok",
        "qdrant": qdrant_status,
    }


@app.get("/stats")
def stats():
    idx = _load_index()
    from qdrant_client import QdrantClient
    try:
        c = QdrantClient(path=QDRANT_PATH)
        info = c.get_collection(COLLECTION_NAME)
        point_count = info.points_count
    except Exception:
        point_count = 0

    review_counts = {}
    source_type_counts = {}
    document_type_counts = {}
    conversion_status_counts = {}
    ingest_ready_books = 0
    ingest_ready_pages = 0
    pending_review_books = 0
    rejected_books = 0
    for record in idx["files"]:
        review_status = str(record.get("review_status", "approved_auto") or "approved_auto")
        review_counts[review_status] = review_counts.get(review_status, 0) + 1
        source_type = infer_source_type(record)
        source_type_counts[source_type] = source_type_counts.get(source_type, 0) + 1
        document_type = infer_document_type(record)
        document_type_counts[document_type] = document_type_counts.get(document_type, 0) + 1
        conversion_status = infer_conversion_status(record)
        conversion_status_counts[conversion_status] = conversion_status_counts.get(conversion_status, 0) + 1
        pages = int(record.get("total_pages", 0) or 0)
        if bool(record.get("ingest_ready", True)):
            ingest_ready_books += 1
            ingest_ready_pages += pages
        if review_status == "pending_review":
            pending_review_books += 1
        elif review_status == "rejected":
            rejected_books += 1

    return {
        "total_books": idx["total_files"],
        "ingest_ready_books": ingest_ready_books,
        "pending_review_books": pending_review_books,
        "rejected_books": rejected_books,
        "total_pages": sum(r["total_pages"] for r in idx["files"]),
        "ingest_ready_pages": ingest_ready_pages,
        "total_points_indexed": point_count,
        "languages": idx["languages"],
        "review_status_counts": review_counts,
        "source_type_counts": source_type_counts,
        "document_type_counts": document_type_counts,
        "conversion_status_counts": conversion_status_counts,
    }


@app.post("/ask", response_model=AskResponse)
def ask(req: AskRequest):
    context_chunks = retrieve(
        query=req.query,
        top_k=req.top_k,
        language=req.language,
        mode="balanced",
    )

    answer, backend_used, mode_used, _chunks = generate(
        query=req.query,
        context_chunks=context_chunks,
        strict=req.strict,
        mode=req.mode,
    )

    sources = extract_sources(answer, context_chunks)

    return AskResponse(
        answer=answer,
        backend_used=backend_used,
        mode=mode_used,
        sources=sources,
    )


@app.get("/search")
def search(q: str, top_k: int = 5, language: str = "id", mode: str = "fast"):
    results = retrieve(query=q, top_k=top_k, language=language, mode=mode)
    return {
        "query": q,
        "mode": mode,
        "results": [
            {
                "text": r["text"],
                "score": r["score"],
                "payload": r["payload"],
                "score_components": r.get("score_components", {}),
                "matched_concepts": r.get("matched_concepts", []),
                "bm25_terms": r.get("bm25_terms", []),
            }
            for r in results
        ],
    }


@app.get("/books")
def list_books(page: int | None = None, limit: int = 100):
    items = [{
        "book_id": r.get("book_id") or os.path.splitext(r["filename"])[0],
        "filename": r["filename"],
        "json_filename": os.path.basename(resolve_index_json_path(r, JSON_DIR)),
        "json_path": r.get("json_path", ""),
        "source_path": r.get("source_path", ""),
        "source_relpath": r.get("source_relpath", ""),
        "source_ext": infer_source_ext(r),
        "source_type": infer_source_type(r),
        "document_type": infer_document_type(r),
        "conversion_status": infer_conversion_status(r),
        "title": r["title"],
        "language": r["language"],
        "total_pages": r["total_pages"],
        "quality_status": r.get("quality_status", "ok"),
        "review_status": r.get("review_status", "approved_auto"),
        "ingest_ready": bool(r.get("ingest_ready", True)),
    } for r in _load_index()["files"]]

    if page is None:
        return items

    limit = max(1, min(int(limit or 100), 500))
    page = max(1, int(page or 1))
    total = len(items)
    total_pages = max(1, (total + limit - 1) // limit)
    page = min(page, total_pages)
    start = (page - 1) * limit
    end = start + limit
    return {
        "items": items[start:end],
        "page": page,
        "limit": limit,
        "total": total,
        "total_pages": total_pages,
    }


@app.get("/library", response_class=HTMLResponse)
def library(page: int = 1, limit: int = 24, q: str = "", lang: str = "", status: str = ""):
    return HTMLResponse(_render_library_html(page=page, limit=limit, q=q, lang=lang, status=status))


@app.get("/books/{book_id}")
def book_detail(book_id: str):
    record = _book_entry(book_id)
    if not record:
        return JSONResponse({"error": "not found"}, status_code=404)

    json_path = _book_json_path(record)
    if not os.path.exists(json_path):
        return JSONResponse({"error": "not found"}, status_code=404)

    with open(json_path, "r", encoding="utf-8") as f:
        book = json.load(f)

    return {
        "book_id": record.get("book_id") or book_id,
        "filename": book["filename"],
        "json_filename": os.path.basename(json_path),
        "json_path": record.get("json_path", ""),
        "source_path": record.get("source_path", ""),
        "source_relpath": record.get("source_relpath", ""),
        "source_ext": infer_source_ext(record),
        "source_type": infer_source_type(record),
        "document_type": infer_document_type(record),
        "conversion_status": infer_conversion_status(record),
        "title": book["title"],
        "language": book["language"],
        "total_pages": book["total_pages"],
        "quality_status": book.get("quality_status", record.get("quality_status", "ok")),
        "review_status": book.get("review_status", record.get("review_status", "approved_auto")),
        "ingest_ready": bool(book.get("ingest_ready", record.get("ingest_ready", True))),
        "pages": [p["page"] for p in book["pages"]],
    }


@app.get("/books/{book_id}/raw")
def book_raw(book_id: str):
    loaded = _load_book_record(book_id)
    if not loaded:
        return JSONResponse({"error": "not found"}, status_code=404)
    return JSONResponse(content=loaded["book"])


@app.get("/books/{book_id}/edit", response_class=HTMLResponse)
def book_edit(book_id: str):
    loaded = _load_book_record(book_id)
    if not loaded:
        return HTMLResponse("<h1>Not found</h1>", status_code=404)
    return HTMLResponse(_render_editor_html(book_id, loaded["record"], loaded["book"]))


@app.post("/books/{book_id}/edit")
def book_edit_save(book_id: str, req: JsonEditRequest):
    loaded = _load_book_record(book_id)
    if not loaded:
        raise HTTPException(status_code=404, detail="not found")

    try:
        edited = json.loads(req.json_text)
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=400, detail=f"invalid JSON: {exc.msg}") from exc

    if not isinstance(edited, dict):
        raise HTTPException(status_code=400, detail="JSON root must be an object")

    if not isinstance(edited.get("pages", []), list):
        raise HTTPException(status_code=400, detail="JSON must contain a pages list")

    record = loaded["record"]
    json_path = loaded["json_path"]

    merged_book = _apply_record_to_book_json(edited, record, json_path)
    now = datetime.now(timezone.utc).isoformat()
    merged_book["review_status"] = "pending_review"
    merged_book["review_required"] = True
    merged_book["review_route"] = "manual_or_lease_coordinator"
    merged_book["reviewed_by"] = "web_editor"
    merged_book["reviewed_at"] = now
    merged_book["review_note"] = "edited via web editor"
    merged_book["ingest_ready"] = False

    _save_json_file(json_path, merged_book)
    extra_record_updates = {
        "quality_status": merged_book.get("quality_status", record.get("quality_status", "ok")),
        "quality_reasons": merged_book.get("quality_reasons", record.get("quality_reasons", [])),
        "quality_warnings": merged_book.get("quality_warnings", record.get("quality_warnings", [])),
        "review_status": "pending_review",
        "review_required": True,
        "review_route": "manual_or_lease_coordinator",
        "reviewed_by": "web_editor",
        "reviewed_at": now,
        "review_note": "edited via web editor",
        "ingest_ready": False,
    }
    updated_record = _update_book_after_json_edit(
        book_id,
        record,
        merged_book,
        json_path,
        release_indexed=True,
        extra_record_updates=extra_record_updates,
    )
    updated_record.update(extra_record_updates)

    return {
        "ok": True,
        "book_id": book_id,
        "json_path": merged_book.get("json_path", os.path.relpath(json_path, JSON_DIR).replace("\\", "/")),
        "review_status": merged_book["review_status"],
        "ingest_ready": merged_book["ingest_ready"],
    }


@app.post("/books/{book_id}/review")
def book_review_action(book_id: str, req: ReviewActionRequest):
    loaded = _load_book_record(book_id)
    if not loaded:
        raise HTTPException(status_code=404, detail="not found")

    action = str(req.action or "").strip().lower()
    scope = str(req.scope or "book").strip().lower()
    record = loaded["record"]
    book = loaded["book"]
    json_path = loaded["json_path"]
    now = datetime.now(timezone.utc).isoformat()

    if scope == "page":
        page_num = req.page_num
        if page_num is None:
            raise HTTPException(status_code=400, detail="page_num is required for page review")
        if action not in {"page_reviewed", "reviewed", "page_pending", "pending_review"}:
            raise HTTPException(status_code=400, detail="invalid page action")
        updated_book = _update_page_review(
            book,
            page_num=page_num,
            action=action,
            reviewed_by=req.reviewed_by,
            note=req.note,
            promote_book=req.promote_book,
        )
        _save_json_file(json_path, updated_book)

        index = _load_index()
        idx_record = _find_index_record(index, book_id)
        if idx_record:
            updated_record = _apply_book_to_record(idx_record, updated_book, json_path)
            updated_record["page_review_status"] = updated_book.get("page_reviews", {}).get(str(page_num), {}).get("status", "")
            updated_record["page_reviewed_by"] = req.reviewed_by
            updated_record["page_reviewed_at"] = now if action in {"page_reviewed", "reviewed"} else ""
            updated_record["page_review_note"] = req.note
            for i, existing in enumerate(index.get("files", [])):
                existing_book_id = existing.get("book_id") or os.path.splitext(existing.get("filename", ""))[0]
                if existing_book_id == book_id:
                    index["files"][i] = updated_record
                    break
            _rebuild_index_summary(index)
            _save_json_file(_index_path(), index)

            content_index = _load_json_file(_content_index_path(), {"entries": [], "total_files": 0})
            if isinstance(content_index, dict):
                content_entry = _build_content_index_entry(updated_book, updated_record, json_path)
                content_index = _replace_content_index_entry(content_index, content_entry)
                _save_json_file(_content_index_path(), content_index)

        return {
            "ok": True,
            "scope": "page",
            "book_id": book_id,
            "page_num": page_num,
            "page_review_status": updated_book.get("page_reviews", {}).get(str(page_num), {}).get("status", ""),
            "review_status": updated_book.get("review_status", book.get("review_status", "approved_auto")),
            "ingest_ready": bool(updated_book.get("ingest_ready", book.get("ingest_ready", True))),
        }

    if action in {"delete", "remove"}:
        delete_result = _delete_book_everywhere(book_id)
        return {"ok": True, "scope": "book", "action": "delete", **delete_result}

    if action in {"duplicate", "mark_duplicate"}:
        note = req.note or "duplicate"
        updated_book = _apply_review_to_book_json(book, record, "rejected", req.reviewed_by, note)
    elif action in {"approved_manual", "approved_lease", "rejected", "pending_review"}:
        updated_book = _apply_review_to_book_json(book, record, action, req.reviewed_by, req.note)
    else:
        raise HTTPException(status_code=400, detail="invalid action")

    _save_json_file(json_path, updated_book)

    index = _load_index()
    idx_record = _find_index_record(index, book_id)
    if not idx_record:
        raise HTTPException(status_code=404, detail="index record not found")

    updated_record = _apply_book_to_record(idx_record, updated_book, json_path)
    updated_record["review_status"] = updated_book.get("review_status", "approved_auto")
    updated_record["review_required"] = bool(updated_book.get("review_required", False))
    updated_record["review_route"] = updated_book.get("review_route", "auto")
    updated_record["reviewed_by"] = updated_book.get("reviewed_by", "")
    updated_record["reviewed_at"] = updated_book.get("reviewed_at", "")
    updated_record["review_note"] = updated_book.get("review_note", "")
    updated_record["ingest_ready"] = bool(updated_book.get("ingest_ready", True))

    for i, existing in enumerate(index.get("files", [])):
        existing_book_id = existing.get("book_id") or os.path.splitext(existing.get("filename", ""))[0]
        if existing_book_id == book_id:
            index["files"][i] = updated_record
            break
    _rebuild_index_summary(index)
    _save_json_file(_index_path(), index)

    content_index = _load_json_file(_content_index_path(), {"entries": [], "total_files": 0})
    if isinstance(content_index, dict):
        content_entry = _build_content_index_entry(updated_book, updated_record, json_path)
        content_index = _replace_content_index_entry(content_index, content_entry)
        _save_json_file(_content_index_path(), content_index)

    if action in {"approved_manual", "approved_lease"}:
        _refresh_lexical_cache()

    return {
        "ok": True,
        "scope": "book",
        "action": action,
        "book_id": book_id,
        "review_status": updated_book.get("review_status", "approved_auto"),
        "ingest_ready": bool(updated_book.get("ingest_ready", True)),
        "json_path": updated_book.get("json_path", os.path.relpath(json_path, JSON_DIR).replace("\\", "/")),
    }


@app.get("/books/{book_id}/pages/{page_num}")
def page_content(book_id: str, page_num: int):
    loaded = _load_book_record(book_id)
    if not loaded:
        return JSONResponse({"error": "not found"}, status_code=404)
    record = loaded["record"]
    book = loaded["book"]
    for p in book["pages"]:
        if p["page"] == page_num:
            return {
                "book_id": book_id,
                "title": book["title"],
                "page": p["page"],
                "content": p["content"],
            }

    return JSONResponse({"error": "page not found"}, status_code=404)


@app.get("/books/{book_id}/pages/{page_num}/view", response_class=HTMLResponse)
def page_content_view(book_id: str, page_num: int, theme: str = "dark", font: int = 19, q: str = ""):
    loaded = _load_book_record(book_id)
    if not loaded:
        return HTMLResponse("<h1>Not found</h1>", status_code=404)

    record = loaded["record"]
    book = loaded["book"]
    page = _find_book_page(book, page_num)
    if not page:
        return HTMLResponse("<h1>Page not found</h1>", status_code=404)

    return HTMLResponse(
        _render_page_html(
            record,
            book,
            page,
            page_num,
            theme=theme,
            font_size=font,
            query=q,
        )
    )


@app.get("/sources/{book_id}/pages/{page_num}", response_class=HTMLResponse)
def source_page_view(book_id: str, page_num: int, theme: str = "dark", font: int = 18, q: str = ""):
    loaded = _load_book_record(book_id)
    if not loaded:
        return HTMLResponse("<h1>Not found</h1>", status_code=404)

    record = loaded["record"]
    book = loaded["book"]
    page = _find_book_page(book, page_num)
    if not page:
        return HTMLResponse("<h1>Page not found</h1>", status_code=404)

    return HTMLResponse(
        _render_source_preview_html(
            record,
            book,
            page_num,
            theme=theme,
            font_size=font,
            query=q,
        )
    )


@app.get("/books/{book_id}/pages/{page_num}/review", response_class=HTMLResponse)
@app.get("/admin/books/{book_id}/pages/{page_num}/review", response_class=HTMLResponse)
def page_review_workspace(book_id: str, page_num: int, theme: str = "dark", font: int = 18, q: str = ""):
    loaded = _load_book_record(book_id)
    if not loaded:
        return HTMLResponse("<h1>Not found</h1>", status_code=404)

    record = loaded["record"]
    book = loaded["book"]
    page = _find_book_page(book, page_num)
    if not page:
        return HTMLResponse("<h1>Page not found</h1>", status_code=404)

    return HTMLResponse(
        _render_page_review_html(
            record,
            book,
            page,
            page_num,
            theme=theme,
            font_size=font,
            q=q,
        )
    )


@app.post("/books/{book_id}/pages/{page_num}/edit")
@app.post("/admin/books/{book_id}/pages/{page_num}/edit")
def page_edit_save(book_id: str, page_num: int, req: PageEditRequest):
    loaded = _load_book_record(book_id)
    if not loaded:
        raise HTTPException(status_code=404, detail="not found")

    record = loaded["record"]
    book = loaded["book"]
    json_path = loaded["json_path"]
    page = _find_book_page(book, page_num)
    if not page:
        raise HTTPException(status_code=404, detail="page not found")

    now = datetime.now(timezone.utc).isoformat()
    page["content"] = req.content
    page["page_review_status"] = "pending_review"
    page["page_reviewed_by"] = req.reviewed_by or "web"
    page["page_reviewed_at"] = ""
    page["page_review_note"] = req.note or "edited via web"

    page_reviews = dict(book.get("page_reviews", {}))
    page_reviews[str(page_num)] = {
        "page": int(page_num),
        "status": "pending_review",
        "reviewed_by": req.reviewed_by or "web",
        "reviewed_at": "",
        "note": req.note or "edited via web",
    }
    book["page_reviews"] = page_reviews
    book["review_status"] = "pending_review"
    book["review_required"] = True
    book["review_route"] = "manual_or_lease_coordinator"
    book["reviewed_by"] = req.reviewed_by or "web"
    book["reviewed_at"] = now
    book["review_note"] = req.note or "edited via web"
    book["ingest_ready"] = False

    _save_json_file(json_path, book)

    extra_record_updates = {
        "page_review_status": "pending_review",
        "page_reviewed_by": req.reviewed_by or "web",
        "page_reviewed_at": "",
        "page_review_note": req.note or "edited via web",
        "review_status": "pending_review",
        "review_required": True,
        "review_route": "manual_or_lease_coordinator",
        "reviewed_by": req.reviewed_by or "web",
        "reviewed_at": now,
        "review_note": req.note or "edited via web",
        "ingest_ready": False,
    }
    updated_record = _update_book_after_json_edit(
        book_id,
        record,
        book,
        json_path,
        release_indexed=True,
        extra_record_updates=extra_record_updates,
    )
    updated_record.update(extra_record_updates)
    index = _load_index()
    for i, existing in enumerate(index.get("files", [])):
        existing_book_id = existing.get("book_id") or os.path.splitext(existing.get("filename", ""))[0]
        if existing_book_id == book_id:
            index["files"][i] = updated_record
            break
    _rebuild_index_summary(index)
    _save_json_file(_index_path(), index)

    return {
        "ok": True,
        "book_id": book_id,
        "page_num": page_num,
        "review_status": book.get("review_status", "pending_review"),
        "page_review_status": page["page_review_status"],
        "ingest_ready": bool(book.get("ingest_ready", False)),
    }


@app.post("/books/{book_id}/pages/{page_num}/repair")
@app.post("/admin/books/{book_id}/pages/{page_num}/repair")
def page_repair_draft(book_id: str, page_num: int, req: PageRepairRequest):
    loaded = _load_book_record(book_id)
    if not loaded:
        raise HTTPException(status_code=404, detail="not found")

    record = loaded["record"]
    book = loaded["book"]
    page = _find_book_page(book, page_num)
    if not page:
        raise HTTPException(status_code=404, detail="page not found")

    source_text, source_page_num, _, source_type, _source_path = _source_page_text(record, page_num)
    current_text = page.get("content", "")
    backend = str(req.backend or "local").strip().lower()
    instruction = " ".join(part for part in [req.instruction, req.note] if part).strip()
    prompt = f"""Anda adalah editor teks untuk buku agama.
Tugas: perbaiki teks hasil ekstraksi/konversi pada halaman berikut tanpa mengubah makna.
Jika teks sumber tidak cukup, pertahankan teks saat ini.
Jangan menambah penjelasan.
Kembalikan hanya teks final.

Metadata:
- Book ID: {record.get('book_id') or book_id}
- Title: {book.get('title', record.get('title', ''))}
- Halaman JSON: {page_num}
- Halaman sumber: {source_page_num}
- Source type: {source_type}
- Instruction: {instruction or '-'}

Teks sumber:
{source_text[:8000]}

Teks saat ini:
{current_text[:8000]}
"""

    try:
        if backend == "lease":
            draft = generate_remote(prompt)
            backend_used = "lease_coordinator"
        elif backend == "large":
            draft = generate_remote(prompt)
            backend_used = "lease_coordinator"
        else:
            draft = generate_local(prompt)
            backend_used = "ollama"
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"repair draft failed: {exc}") from exc

    return {
        "ok": True,
        "book_id": book_id,
        "page_num": page_num,
        "backend_used": backend_used,
        "suggested_content": draft,
        "instruction": instruction,
        "source_page_num": source_page_num,
    }


@app.post("/books/{book_id}/pages/{page_num}/apply-markers")
@app.post("/admin/books/{book_id}/pages/{page_num}/apply-markers")
def page_apply_markers(book_id: str, page_num: int, req: PageMarkerApplyRequest):
    loaded = _load_book_record(book_id)
    if not loaded:
        raise HTTPException(status_code=404, detail="not found")

    record = loaded["record"]
    book = loaded["book"]
    page = _find_book_page(book, page_num)
    if not page:
        raise HTTPException(status_code=404, detail="page not found")

    result = apply_reference_markers(
        req.content,
        dorar_policy=req.dorar_policy,
        dorar_choices=req.dorar_choices,
    )
    diff_text = _build_unified_diff(req.content, result["resolved_text"], from_label="draft", to_label="resolved")
    diff_summary = "Tidak ada perubahan."
    if result["resolved_text"] != req.content:
        before_lines = len((req.content or "").splitlines()) or (1 if req.content else 0)
        after_lines = len((result["resolved_text"] or "").splitlines()) or (1 if result["resolved_text"] else 0)
        diff_summary = f"Perubahan terdeteksi: {before_lines} baris → {after_lines} baris."
    return {
        "ok": True,
        "book_id": book_id,
        "page_num": page_num,
        "resolved_content": result["resolved_text"],
        "diff_text": diff_text,
        "diff_summary": diff_summary,
        "stages": result["stages"],
        "unresolved": result["unresolved"],
        "dorar_candidates": result["dorar_candidates"],
        "quran_reference_path": result["quran_reference_path"],
        "hadith_reference_dir": result["hadith_reference_dir"],
        "source_path": record.get("source_path", ""),
    }


@app.post("/admin/hadith/dorar/search")
def dorar_search(req: DorarSearchRequest):
    query = (req.query or "").strip()
    if not query:
        raise HTTPException(status_code=400, detail="query is required")
    limit = max(1, min(int(req.limit or 5), 10))
    results = search_dorar_candidates(query, limit=limit)
    return {
        "ok": True,
        "query": query,
        "limit": limit,
        "results": results,
    }


@app.post("/admin/hadith/local/search")
def local_hadith_search(req: LocalHadithSearchRequest):
    query = (req.query or "").strip()
    if not query:
        raise HTTPException(status_code=400, detail="query is required")
    limit = max(1, min(int(req.limit or 10), 50))
    collection = (req.collection or "").strip()
    results = search_local_hadith(query, limit=limit, collection=collection)
    return {
        "ok": True,
        "query": query,
        "collection": collection,
        "limit": limit,
        "results": results,
    }


@app.post("/debug/retrieve")
def debug_retrieve(req: RetrieveRequest):
    results = retrieve(query=req.query, top_k=req.top_k, language=req.language, mode=req.mode)
    return {
        "query": req.query,
        "mode": req.mode,
        "results": [
            {
                "text": r["text"],
                "score": r["score"],
                "payload": r["payload"],
                "score_components": r.get("score_components", {}),
                "matched_concepts": r.get("matched_concepts", []),
                "bm25_terms": r.get("bm25_terms", []),
            }
            for r in results
        ],
    }


STATIC_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static")
if os.path.exists(STATIC_DIR):
    app.mount("/", StaticFiles(directory=STATIC_DIR, html=True), name="static")
