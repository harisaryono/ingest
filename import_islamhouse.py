#!/usr/bin/env python3
"""Import the Islamhouse corpus into json_output with content dedupe.

This script scans a source directory, extracts text from supported formats,
checks whether the extracted content already exists in the current JSON corpus,
marks suspicious extractions as pending review, and writes only unique books
into DATABASE_DIR/json_output.

Supported inputs:
- .txt
- .html / .htm
- .docx
- .doc (via LibreOffice HTML fallback)
- .pdf
- .epub / .ibooks (ZIP-based XHTML extraction)

Skipped inputs are reported to JSONL audit files so the caller can audit what
was already present in the corpus or what failed quality checks.
"""

from __future__ import annotations

import argparse
import hashlib
import html as html_lib
import json
import os
import re
import shutil
import subprocess
import sys
import signal
import tempfile
import unicodedata
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Tuple

from bs4 import BeautifulSoup
from docx import Document
from pdfminer.high_level import extract_text as pdf_extract_text

SCRIPT_DIR = Path(__file__).resolve().parent
DEFAULT_INPUT_DIR = Path("/media/harry/DATA250/Islamhouse")
DATABASE_DIR = Path(
    os.getenv(
        "DATABASE_DIR",
        str(Path(__file__).resolve().parent.parent.parent / "DATABASE"),
    )
)
OUTPUT_DIR = DATABASE_DIR / "json_output"
CONTENT_INDEX_PATH = OUTPUT_DIR / "_content_index.json"
DUPLICATE_REPORT_PATH = OUTPUT_DIR / "_duplicates.jsonl"
QUALITY_REPORT_PATH = OUTPUT_DIR / "_quality_issues.jsonl"
INDEX_PATH = OUTPUT_DIR / "_index.json"

SUPPORTED_EXTS = {
    ".txt",
    ".htm",
    ".html",
    ".docx",
    ".doc",
    ".pdf",
    ".epub",
    ".ibooks",
}


def source_type_from_path(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in {".htm", ".html"}:
        return "html"
    if ext:
        return ext.lstrip(".")
    return "unknown"


def conversion_status_from_quality(quality_status: str) -> str:
    quality_status = (quality_status or "").strip().lower()
    if quality_status == "quarantine":
        return "failed"
    if quality_status == "warn":
        return "degraded"
    if quality_status == "ok":
        return "good"
    return "unknown"

ARABIC_DIACRITICS_RE = re.compile(r"[\u064b-\u065f\u0670\u06d6-\u06ed]")
WORD_RE = re.compile(r"[A-Za-z0-9\u0600-\u06ff]+")
LANG_PREFIX_RE = re.compile(r"^([a-z]{2})(?:[_-]|$)", re.IGNORECASE)
PAGE_BREAK_RE = re.compile(r"\f+")
ARABIC_RANGES = (
    (0x0600, 0x06FF),
    (0x0750, 0x077F),
    (0x08A0, 0x08FF),
    (0xFB50, 0xFDFF),
    (0xFE70, 0xFEFF),
)


@dataclass
class BookSignature:
    page_hashes: List[str]
    content_hash: str
    text_hash: str
    text_simhash: int
    page_count: int
    size_bytes: int


def log(message: str) -> None:
    print(message, flush=True)


def sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def is_arabic_char(ch: str) -> bool:
    if not ch:
        return False
    codepoint = ord(ch)
    return any(start <= codepoint <= end for start, end in ARABIC_RANGES)


def detect_script(text: str) -> str:
    arabic_chars = 0
    latin_chars = 0
    for ch in text:
        if is_arabic_char(ch):
            arabic_chars += 1
        elif "LATIN" in unicodedata.name(ch, ""):
            latin_chars += 1

    if arabic_chars == 0 and latin_chars == 0:
        return "unknown"
    if arabic_chars >= max(20, latin_chars * 2):
        return "arabic"
    if latin_chars >= max(20, arabic_chars * 2):
        return "latin"
    return "mixed"


def char_profile(text: str) -> Dict[str, int]:
    profile = {
        "total_chars": len(text),
        "nonspace_chars": 0,
        "letters": 0,
        "arabic_chars": 0,
        "latin_chars": 0,
        "digits": 0,
        "punct": 0,
        "symbols": 0,
        "other": 0,
        "replacement": text.count("\ufffd"),
    }

    for ch in text:
        if not ch.isspace():
            profile["nonspace_chars"] += 1
        category = unicodedata.category(ch)
        if category.startswith("L"):
            profile["letters"] += 1
            if is_arabic_char(ch):
                profile["arabic_chars"] += 1
            elif "LATIN" in unicodedata.name(ch, ""):
                profile["latin_chars"] += 1
        elif category.startswith("N"):
            profile["digits"] += 1
        elif category.startswith("P"):
            profile["punct"] += 1
        elif category.startswith("S"):
            profile["symbols"] += 1
        else:
            profile["other"] += 1

    return profile


def write_jsonl(path: Path, entry: Dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as f:
        f.write(json.dumps(entry, ensure_ascii=False) + "\n")


class FileTimeoutError(TimeoutError):
    pass


def run_with_timeout(timeout_seconds: int, func, *args, **kwargs):
    if timeout_seconds <= 0:
        return func(*args, **kwargs)

    def _alarm_handler(signum, frame):  # noqa: ARG001
        raise FileTimeoutError(f"file processing exceeded {timeout_seconds}s")

    previous_handler = signal.signal(signal.SIGALRM, _alarm_handler)
    signal.setitimer(signal.ITIMER_REAL, timeout_seconds)
    try:
        return func(*args, **kwargs)
    finally:
        signal.setitimer(signal.ITIMER_REAL, 0)
        signal.signal(signal.SIGALRM, previous_handler)


def simhash(tokens: Iterable[str]) -> int:
    counts: Dict[str, int] = {}
    for token in tokens:
        counts[token] = counts.get(token, 0) + 1
    if not counts:
        return 0

    vector = [0] * 64
    for token, weight in counts.items():
        digest = hashlib.blake2b(token.encode("utf-8"), digest_size=8).digest()
        fingerprint = int.from_bytes(digest, "big")
        for bit in range(64):
            if fingerprint & (1 << bit):
                vector[bit] += weight
            else:
                vector[bit] -= weight

    value = 0
    for bit, score in enumerate(vector):
        if score > 0:
            value |= 1 << bit
    return value


def hamming_distance(a: int, b: int) -> int:
    return (a ^ b).bit_count()


def resolve_index_json_path(record: Dict, json_dir: str) -> str:
    json_path = record.get("json_path", "")
    if json_path:
        if os.path.isabs(json_path):
            return json_path
        normalized = os.path.normpath(json_path)
        prefix = f"json_output{os.sep}"
        if normalized == "json_output":
            return json_dir
        if normalized.startswith(prefix):
            normalized = normalized[len(prefix) :]
        return os.path.join(json_dir, normalized)
    filename = record.get("filename", "")
    return os.path.join(json_dir, f"{filename}.json")


def detect_language(source: Path | str) -> str:
    path = source if isinstance(source, Path) else Path(str(source))
    match = LANG_PREFIX_RE.match(path.name)
    if match:
        return match.group(1).lower()
    for part in reversed(path.parts):
        if part.lower() in {"ar", "en", "id", "ru"}:
            return part.lower()
    return "unknown"


def clean_title(filename: str) -> str:
    name = filename
    match = LANG_PREFIX_RE.match(name)
    if match:
        name = name[match.end() :]
    name = name.replace("_", " ").replace("-", " ")
    name = re.sub(r"\s+", " ", name).strip()
    words = []
    for word in name.split():
        if word.upper() == word and len(word) > 1:
            words.append(word)
        elif word.lower() in {"and", "the", "of", "in", "to", "for", "a", "an", "or", "by", "is", "on", "its"}:
            words.append(word.lower() if words else word.capitalize())
        elif word.startswith("al-") or word.startswith("Al-"):
            words.append(word[:1].upper() + word[1:])
        else:
            words.append(word.capitalize())
    if words:
        words[0] = words[0].capitalize()
    return " ".join(words) if words else filename


def make_book_id(source_label: str, relpath: str) -> str:
    normalized = relpath.replace("/", "__").replace("\\", "__")
    normalized = re.sub(r"[^A-Za-z0-9._-]+", "_", normalized)
    normalized = re.sub(r"_+", "_", normalized).strip("_")
    return f"{source_label}__{normalized}"


def normalize_text(text: str, language: str) -> str:
    text = html_lib.unescape(text or "")
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("\x0c", " ").replace("\xa0", " ")
    if language == "ar" or re.search(r"[\u0600-\u06ff]", text):
        text = ARABIC_DIACRITICS_RE.sub("", text)
    text = text.lower()
    text = re.sub(r"[^\w\s\u0600-\u06ff]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def split_pages_from_text(text: str) -> List[str]:
    if not text or not text.strip():
        return []
    pages = PAGE_BREAK_RE.split(text)
    pages = [p.strip() for p in pages if p and p.strip()]
    return pages


def read_text_file(path: Path) -> List[str]:
    with path.open("r", encoding="utf-8", errors="replace") as f:
        content = f.read()
    pages = split_pages_from_text(content)
    return pages or [content.strip()]


def html_to_pages(html_text: str) -> List[str]:
    soup = BeautifulSoup(html_text, "lxml")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    text = soup.get_text("\n")
    pages = split_pages_from_text(text)
    return pages or [text.strip()]


def read_html_file(path: Path) -> List[str]:
    return html_to_pages(path.read_text(encoding="utf-8", errors="replace"))


def read_docx_file(path: Path) -> List[str]:
    doc = Document(str(path))
    pages: List[str] = []
    current: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            current.append(text)
        para_xml = para._p.xml
        if ("w:type=\"page\"" in para_xml or "w:lastRenderedPageBreak" in para_xml) and current:
            pages.append("\n".join(current).strip())
            current = []

    if current:
        pages.append("\n".join(current).strip())

    if not pages:
        all_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        return split_pages_from_text(all_text) or [all_text.strip()]

    return pages


def read_pdf_file(path: Path) -> List[str]:
    text = pdf_extract_text(str(path))
    pages = split_pages_from_text(text)
    return pages or [text.strip()]


def read_zip_xhtml_file(path: Path) -> List[str]:
    pages: List[str] = []
    with zipfile.ZipFile(path) as zf:
        members = sorted(
            name for name in zf.namelist()
            if name.lower().endswith((".xhtml", ".html", ".htm"))
        )
        for name in members:
            with zf.open(name) as handle:
                raw = handle.read().decode("utf-8", errors="replace")
            page_text = html_to_pages(raw)
            pages.extend(page_text or [""])
    return [p for p in pages if p.strip()]


def read_doc_with_libreoffice(path: Path, timeout_seconds: int | None = None) -> List[str]:
    if shutil.which("libreoffice") is None:
        raise RuntimeError("libreoffice not available for .doc fallback")

    with tempfile.TemporaryDirectory(prefix="islamhouse-lo-") as tmp:
        tmpdir = Path(tmp)
        try:
            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to",
                    "html",
                    "--outdir",
                    str(tmpdir),
                    str(path),
                ],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                check=False,
                timeout=timeout_seconds,
            )
        except subprocess.TimeoutExpired as e:
            raise FileTimeoutError(f"LibreOffice conversion timed out for {path.name}") from e
        html_files = sorted(tmpdir.glob("*.html"))
        if result.returncode != 0 or not html_files:
            raise RuntimeError(
                f"LibreOffice conversion failed for {path.name}: {result.stderr.strip() or result.stdout.strip()}"
            )
        html_text = html_files[0].read_text(encoding="utf-8", errors="replace")
        pages = html_to_pages(html_text)
        return pages or [html_text.strip()]


def analyze_quality(pages: List[str], language: str, path: Path, extractor: str, size_bytes: int) -> Dict:
    raw_pages = [page or "" for page in pages]
    normalized_pages = [unicodedata.normalize("NFKC", html_lib.unescape(page)).replace("\x0c", " ").replace("\xa0", " ") for page in raw_pages]
    joined = "\n".join(normalized_pages).strip()
    profile = char_profile(joined)
    tokens = WORD_RE.findall(normalize_text(joined, language))
    unique_tokens = set(tokens)
    page_lengths = [len(page.strip()) for page in raw_pages]
    page_hashes = [sha256_text(normalize_text(page, language)) for page in raw_pages if normalize_text(page, language)]
    page_hash_count = len(page_hashes)
    duplicate_pages = max(0, page_hash_count - len(set(page_hashes)))

    token_count = len(tokens)
    unique_word_ratio = len(unique_tokens) / token_count if token_count else 0.0
    avg_page_chars = sum(page_lengths) / max(len(page_lengths), 1)
    short_page_ratio = sum(1 for length in page_lengths if length < 40) / max(len(page_lengths), 1)
    duplicate_page_ratio = duplicate_pages / max(page_hash_count, 1)
    symbol_ratio = (profile["punct"] + profile["symbols"]) / max(profile["letters"] + profile["digits"] + profile["punct"] + profile["symbols"], 1)
    replacement_ratio = profile["replacement"] / max(profile["total_chars"], 1)
    arabic_ratio = profile["arabic_chars"] / max(profile["letters"], 1)
    latin_ratio = profile["latin_chars"] / max(profile["letters"], 1)
    unique_char_ratio = len(set(joined)) / max(profile["nonspace_chars"], 1)
    detected_script = detect_script(joined)
    expected_language = detect_language(path)
    expected_arabic = expected_language == "ar" or language == "ar" or any(part.lower() == "ar" for part in path.parts)

    reasons: List[str] = []
    warnings: List[str] = []

    if profile["nonspace_chars"] < 40 or token_count < 8:
        reasons.append("too_short_or_empty")
    if profile["nonspace_chars"] >= 120 and unique_char_ratio < 0.04:
        reasons.append("low_character_diversity")
    if replacement_ratio >= 0.01:
        reasons.append("encoding_replacement_chars")
    if duplicate_page_ratio >= 0.67 and page_hash_count >= 3:
        reasons.append("repeated_pages")

    if expected_arabic:
        if profile["arabic_chars"] == 0 and profile["latin_chars"] >= 20 and profile["nonspace_chars"] >= 40:
            reasons.append("arabic_script_mismatch")
        elif profile["arabic_chars"] < 10 and profile["nonspace_chars"] >= 60:
            reasons.append("arabic_missing")
        elif arabic_ratio < 0.12 and profile["nonspace_chars"] >= 200:
            reasons.append("arabic_low_density")
        elif detected_script == "latin" and profile["nonspace_chars"] >= 40:
            reasons.append("arabic_script_mismatch")

    if token_count >= 120 and unique_word_ratio < 0.18:
        warnings.append("low_vocab_diversity")
    if short_page_ratio >= 0.6 and len(page_lengths) >= 2:
        warnings.append("many_short_pages")
    if symbol_ratio >= 0.45 and profile["nonspace_chars"] >= 120:
        warnings.append("high_symbol_density")
    if not expected_arabic and profile["arabic_chars"] >= 20 and profile["latin_chars"] >= 20 and profile["nonspace_chars"] >= 300:
        warnings.append("mixed_script_content")
    if extractor in {"pdf", "libreoffice-html"} and size_bytes >= 1024 * 1024 and profile["nonspace_chars"] < 400:
        reasons.append("conversion_loss_suspected")

    status = "ok"
    if reasons:
        status = "quarantine"
    elif warnings:
        status = "warn"

    return {
        "status": status,
        "reasons": reasons,
        "warnings": warnings,
        "expected_language": expected_language,
        "expected_arabic": expected_arabic,
        "detected_script": detected_script,
        "metrics": {
            **profile,
            "page_count": len(raw_pages),
            "token_count": token_count,
            "unique_word_ratio": round(unique_word_ratio, 4),
            "avg_page_chars": round(avg_page_chars, 2),
            "short_page_ratio": round(short_page_ratio, 4),
            "duplicate_page_ratio": round(duplicate_page_ratio, 4),
            "symbol_ratio": round(symbol_ratio, 4),
            "replacement_ratio": round(replacement_ratio, 4),
            "arabic_ratio": round(arabic_ratio, 4),
            "latin_ratio": round(latin_ratio, 4),
            "unique_char_ratio": round(unique_char_ratio, 4),
        },
    }


def extract_pages(path: Path, timeout_seconds: int | None = None) -> Tuple[List[str], str]:
    ext = path.suffix.lower()
    if ext == ".txt":
        return read_text_file(path), "txt"
    if ext in {".htm", ".html"}:
        return read_html_file(path), "html"
    if ext == ".docx":
        return read_docx_file(path), "docx"
    if ext == ".pdf":
        return read_pdf_file(path), "pdf"
    if ext in {".epub", ".ibooks"}:
        return read_zip_xhtml_file(path), "zip-xhtml"
    if ext == ".doc":
        return read_doc_with_libreoffice(path, timeout_seconds=timeout_seconds), "libreoffice-html"
    raise RuntimeError(f"Unsupported extension: {ext}")


def build_signature(pages: List[str], language: str, size_bytes: int) -> BookSignature:
    page_hashes: List[str] = []
    full_normalized_parts: List[str] = []
    for page in pages:
        normalized = normalize_text(page, language)
        if not normalized:
            continue
        page_hashes.append(sha256_text(normalized))
        full_normalized_parts.append(normalized)
    text = " ".join(full_normalized_parts)
    content_hash = sha256_text("\n".join(page_hashes))
    text_hash = sha256_text(text)
    text_simhash = simhash(text.split())
    return BookSignature(
        page_hashes=page_hashes,
        content_hash=content_hash,
        text_hash=text_hash,
        text_simhash=text_simhash,
        page_count=len(page_hashes),
        size_bytes=size_bytes,
    )


def load_json(path: Path, default):
    if not path.exists():
        return default
    try:
        with path.open("r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return default
    return data if isinstance(data, type(default)) else default


def save_json(path: Path, data) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    tmp = path.with_suffix(path.suffix + ".tmp")
    with tmp.open("w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2, sort_keys=True)
    os.replace(tmp, path)


def build_content_catalog(index_records: List[Dict], output_dir: Path) -> List[Dict]:
    catalog: List[Dict] = []
    for record in index_records:
        json_path = resolve_index_json_path(record, str(output_dir))
        path = Path(json_path)
        if not path.exists():
            continue
        try:
            with path.open("r", encoding="utf-8") as f:
                book = json.load(f)
            pages = [p.get("content", "") for p in book.get("pages", [])]
            sig = build_signature(pages, book.get("language", "unknown"), int(book.get("size_bytes", 0) or 0))
            catalog.append({
                "book_id": record.get("book_id") or book.get("book_id") or os.path.splitext(record["filename"])[0],
                "filename": record["filename"],
                "json_path": record["json_path"],
                "language": record.get("language", book.get("language", "unknown")),
                "title": record.get("title", book.get("title", "")),
                "conversion_status": record.get("conversion_status", book.get("conversion_status", "unknown")),
                "quality_status": record.get("quality_status", book.get("quality_status", "ok")),
                "review_status": record.get("review_status", book.get("review_status", "approved_auto")),
                "review_required": bool(record.get("review_required", book.get("review_required", False))),
                "review_route": record.get("review_route", book.get("review_route", "auto")),
                "reviewed_by": record.get("reviewed_by", book.get("reviewed_by", "")),
                "reviewed_at": record.get("reviewed_at", book.get("reviewed_at", "")),
                "review_note": record.get("review_note", book.get("review_note", "")),
                "ingest_ready": bool(record.get("ingest_ready", book.get("ingest_ready", True))),
                "total_pages": int(record.get("total_pages", len(book.get("pages", []))) or 0),
                "size_bytes": int(record.get("size_bytes", book.get("size_bytes", 0)) or 0),
                "content_hash": sig.content_hash,
                "text_hash": sig.text_hash,
                "text_simhash": sig.text_simhash,
                "page_hashes": sig.page_hashes,
                "source_path": record.get("source_path", ""),
                "source_hash": record.get("source_hash", ""),
            })
        except Exception as e:
            log(f"  WARN  content catalog skip {record.get('filename')}: {e}")
    return catalog


def load_catalog(output_dir: Path) -> Tuple[Dict, List[Dict]]:
    index = load_json(INDEX_PATH, {"total_files": 0, "languages": {}, "files": []})
    entries = build_content_catalog(index.get("files", []), output_dir)
    return index, entries


def page_overlap_score(a: List[str], b: List[str]) -> float:
    if not a or not b:
        return 0.0
    sa = set(a)
    sb = set(b)
    inter = len(sa & sb)
    union = len(sa | sb)
    return inter / union if union else 0.0


def find_duplicate(candidate: Dict, catalog: List[Dict], threshold: float) -> Tuple[Optional[Dict], float, str]:
    if not candidate["page_hashes"]:
        return None, 0.0, "empty"
    candidate_set = set(candidate["page_hashes"])
    best: Optional[Dict] = None
    best_score = 0.0
    best_reason = "none"
    candidate_title_tokens = set(normalize_text(candidate.get("title", ""), candidate["language"]).split())

    for entry in catalog:
        if candidate["language"] != "unknown" and entry.get("language") not in {"unknown", candidate["language"]}:
            continue
        if candidate["page_count"] and entry.get("total_pages"):
            ratio = min(candidate["page_count"], int(entry["total_pages"])) / max(candidate["page_count"], int(entry["total_pages"]))
            if ratio < 0.45:
                continue
        existing_set = set(entry.get("page_hashes", []))
        if not existing_set:
            continue
        if candidate["content_hash"] == entry.get("content_hash"):
            return entry, 1.0, "exact"
        if candidate.get("text_hash") and candidate.get("text_hash") == entry.get("text_hash"):
            return entry, 1.0, "text_exact"

        simhash_distance = None
        if candidate.get("text_simhash") is not None and entry.get("text_simhash") is not None:
            simhash_distance = hamming_distance(int(candidate["text_simhash"]), int(entry["text_simhash"]))

        title_tokens = set(normalize_text(entry.get("title", ""), entry.get("language", "unknown")).split())
        title_score = 0.0
        if candidate_title_tokens or title_tokens:
            title_score = len(candidate_title_tokens & title_tokens) / max(len(candidate_title_tokens | title_tokens), 1)

        page_score = page_overlap_score(list(candidate_set), list(existing_set))
        score = page_score
        score_reason = "page_overlap"
        if simhash_distance is not None:
            simhash_score = 1.0 - (simhash_distance / 64.0)
            if simhash_score > score:
                score = simhash_score
                score_reason = "text_similarity"
        if score >= 0.80 and title_score >= 0.40:
            score = min(1.0, score + min(0.08, title_score * 0.08))
            score_reason = f"{score_reason}+title"
        if score > best_score:
            best = entry
            best_score = score
            best_reason = score_reason

    if best and best_score >= threshold:
        return best, best_score, best_reason
    return None, best_score, best_reason


def scan_files(input_dir: Path, recursive: bool) -> List[Path]:
    if recursive:
        paths = [p for p in input_dir.rglob("*") if p.is_file()]
    else:
        paths = [p for p in input_dir.iterdir() if p.is_file()]
    return sorted(
        p for p in paths
        if p.suffix.lower() in SUPPORTED_EXTS
        and not p.name.startswith(".~lock.")
    )


def make_output_path(source_label: str, input_dir: Path, path: Path) -> Tuple[Path, str]:
    relpath = path.relative_to(input_dir).as_posix()
    rel_json = Path(source_label) / f"{relpath}.json"
    outpath = OUTPUT_DIR / rel_json
    return outpath, relpath


def process_file(
    path: Path,
    input_dir: Path,
    source_label: str,
    catalog: List[Dict],
    threshold: float,
    keep_duplicates: bool,
    file_timeout_seconds: int,
) -> Tuple[Optional[Dict], Optional[Dict], Optional[Dict], str]:
    pages, extractor = extract_pages(path, timeout_seconds=file_timeout_seconds)
    if not pages:
        return None, None, None, extractor

    language = detect_language(path)
    signature = build_signature(pages, language, path.stat().st_size)
    outpath, relpath = make_output_path(source_label, input_dir, path)
    book_id = make_book_id(source_label, relpath)
    source_path = str(path.resolve())
    source_relpath = relpath
    source_ext = path.suffix.lower()
    source_type = source_type_from_path(path)
    document_type = "html_document" if source_type == "html" else "book"
    quality = analyze_quality(pages, language, path, extractor, path.stat().st_size)
    duplicate_of, score, reason = find_duplicate(
        {
            "page_hashes": signature.page_hashes,
            "content_hash": signature.content_hash,
            "text_hash": signature.text_hash,
            "text_simhash": signature.text_simhash,
            "page_count": signature.page_count,
            "language": language,
            "title": clean_title(path.stem),
        },
        catalog,
        threshold,
    )

    record = {
        "book_id": book_id,
        "filename": path.name,
        "language": language,
        "title": clean_title(path.stem),
        "size_bytes": path.stat().st_size,
        "total_pages": len(pages),
        "json_path": str(outpath.relative_to(OUTPUT_DIR)),
        "source_root": str(input_dir),
        "source_path": source_path,
        "source_relpath": source_relpath,
        "source_ext": source_ext,
        "source_type": source_type,
        "document_type": document_type,
        "source_hash": sha256_file(path),
        "content_hash": signature.content_hash,
        "text_hash": signature.text_hash,
        "text_simhash": signature.text_simhash,
        "extractor": extractor,
        "quality_status": quality["status"],
        "quality_reasons": quality["reasons"],
        "quality_warnings": quality["warnings"],
        "quality_metrics": quality["metrics"],
        "quality_expected_language": quality["expected_language"],
        "quality_expected_arabic": quality["expected_arabic"],
        "quality_detected_script": quality["detected_script"],
        "conversion_status": conversion_status_from_quality(quality["status"]),
    }

    if quality["status"] == "quarantine":
        record.update(
            {
                "review_status": "pending_review",
                "review_required": True,
                "review_route": "manual_or_lease_coordinator",
                "reviewed_by": "",
                "reviewed_at": "",
                "ingest_ready": False,
            }
        )
    else:
        record.update(
            {
                "review_status": "approved_auto",
                "review_required": False,
                "review_route": "auto",
                "reviewed_by": "",
                "reviewed_at": "",
                "ingest_ready": True,
            }
        )

    book_json = {
        **record,
        "pages": [{"page": i + 1, "content": page} for i, page in enumerate(pages)],
    }

    return book_json, duplicate_of, {"score": score, "reason": reason}, extractor


def update_index(index: Dict, record: Dict) -> None:
    files = index.setdefault("files", [])
    files = [r for r in files if r.get("json_path") != record["json_path"]]
    files.append(record)
    files.sort(key=lambda r: (r.get("language", ""), r.get("json_path", "")))
    index["files"] = files
    index["total_files"] = len(files)
    languages: Dict[str, int] = {}
    source_types: Dict[str, int] = {}
    document_types: Dict[str, int] = {}
    conversion_status_counts: Dict[str, int] = {}
    for r in files:
        lang = r.get("language", "unknown")
        languages[lang] = languages.get(lang, 0) + 1
        source_type = r.get("source_type", "unknown")
        source_types[source_type] = source_types.get(source_type, 0) + 1
        document_type = r.get("document_type", "book")
        document_types[document_type] = document_types.get(document_type, 0) + 1
        conversion_status = r.get("conversion_status", "unknown")
        conversion_status_counts[conversion_status] = conversion_status_counts.get(conversion_status, 0) + 1
    index["languages"] = languages
    index["source_types"] = source_types
    index["document_types"] = document_types
    index["conversion_status_counts"] = conversion_status_counts


def update_content_index(entries: List[Dict], record: Dict, signature: BookSignature) -> List[Dict]:
    entries = [e for e in entries if e.get("json_path") != record["json_path"]]
    entries.append({
        "book_id": record["book_id"],
        "filename": record["filename"],
        "language": record["language"],
        "title": record["title"],
        "json_path": record["json_path"],
        "source_path": record["source_path"],
        "source_relpath": record["source_relpath"],
        "source_ext": record.get("source_ext", ""),
        "source_type": record.get("source_type", "unknown"),
        "document_type": record.get("document_type", "book"),
        "conversion_status": record.get("conversion_status", "unknown"),
        "source_hash": record["source_hash"],
        "quality_status": record.get("quality_status", "ok"),
        "quality_reasons": record.get("quality_reasons", []),
        "quality_warnings": record.get("quality_warnings", []),
        "review_status": record.get("review_status", "approved_auto"),
        "review_required": bool(record.get("review_required", False)),
        "review_route": record.get("review_route", "auto"),
        "content_hash": signature.content_hash,
        "text_hash": signature.text_hash,
        "text_simhash": signature.text_simhash,
        "page_count": signature.page_count,
        "size_bytes": signature.size_bytes,
        "page_hashes": signature.page_hashes,
    })
    entries.sort(key=lambda r: r.get("json_path", ""))
    return entries


def write_duplicate_report(path: Path, entry: Dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as f:
        f.write(json.dumps(entry, ensure_ascii=False) + "\n")


def run_canonicalizer(prune_qdrant: bool = True) -> int:
    script = SCRIPT_DIR / "rag" / "canonicalize_islamhouse.py"
    if not script.exists():
        log(f"Canonicalizer missing: {script}")
        return 1

    cmd = [sys.executable, str(script)]
    if prune_qdrant:
        cmd.append("--prune-qdrant")
    log(f"Running canonicalizer: {' '.join(cmd)}")
    proc = subprocess.run(cmd, cwd=str(SCRIPT_DIR), check=False)
    log(f"Canonicalizer finished with rc={proc.returncode}")
    return proc.returncode


def main() -> None:
    parser = argparse.ArgumentParser(description="Import Islamhouse sources into json_output")
    parser.add_argument("--input-dir", default=str(DEFAULT_INPUT_DIR), help="source directory to scan")
    parser.add_argument("--source-label", default="islamhouse", help="output subdirectory and book_id prefix")
    parser.add_argument("--recursive", action="store_true", help="scan source directory recursively")
    parser.add_argument("--keep-duplicates", action="store_true", help="write duplicate books instead of skipping them")
    parser.add_argument("--keep-suspect", action="store_true", help="legacy flag; suspect books are still written but remain pending review")
    parser.add_argument("--duplicate-threshold", type=float, default=0.92, help="page overlap threshold for duplicate detection")
    parser.add_argument("--limit", type=int, default=0, help="process at most N files for testing")
    parser.add_argument("--file-timeout-seconds", type=int, default=int(os.getenv("IMPORT_FILE_TIMEOUT_SECONDS", "600")), help="skip a file if processing exceeds this many seconds")
    parser.add_argument("--skip-canonicalize", action="store_true", help="skip post-import family canonicalization")
    parser.add_argument("--no-prune-qdrant", action="store_true", help="keep Qdrant state untouched when canonicalizing")
    args = parser.parse_args()

    input_dir = Path(args.input_dir).expanduser().resolve()
    if not input_dir.exists():
        raise SystemExit(f"Input directory not found: {input_dir}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.joinpath(args.source_label).mkdir(parents=True, exist_ok=True)

    index, catalog = load_catalog(OUTPUT_DIR)
    files = scan_files(input_dir, args.recursive)
    if args.limit and args.limit > 0:
        files = files[: args.limit]

    log(f"Source dir       : {input_dir}")
    log(f"Output dir       : {OUTPUT_DIR}")
    log(f"Source label     : {args.source_label}")
    log(f"Files discovered : {len(files)}")
    log(f"Existing books   : {len(index.get('files', []))}")
    log(f"Catalog entries  : {len(catalog)}")
    log(f"File timeout     : {args.file_timeout_seconds}s")

    processed = 0
    imported = 0
    skipped_duplicates = 0
    quarantined = 0
    unsupported = 0
    duplicate_hits = 0
    quality_warned = 0

    for i, path in enumerate(files, 1):
        processed += 1
        try:
            book_json, duplicate_of, dup_info, extractor = run_with_timeout(
                args.file_timeout_seconds,
                process_file,
                path,
                input_dir,
                args.source_label,
                catalog,
                args.duplicate_threshold,
                args.keep_duplicates,
                args.file_timeout_seconds,
            )
        except FileTimeoutError as e:
            unsupported += 1
            write_jsonl(
                QUALITY_REPORT_PATH,
                {
                    "kind": "timeout",
                    "source_path": str(path.resolve()),
                    "source_relpath": path.relative_to(input_dir).as_posix(),
                    "filename": path.name,
                    "source_ext": path.suffix.lower(),
                    "source_type": source_type_from_path(path),
                    "document_type": "html_document" if source_type_from_path(path) == "html" else "book",
                    "status": "skipped",
                    "reason": "file_processing_timeout",
                    "timeout_seconds": args.file_timeout_seconds,
                    "error": str(e),
                },
            )
            log(f"[{i:04d}] TIMEOUT {path.name} skipped after {args.file_timeout_seconds}s")
            continue
        except Exception as e:
            unsupported += 1
            log(f"[{i:04d}] ERROR  {path.name}: {e}")
            continue

        if book_json is None:
            unsupported += 1
            log(f"[{i:04d}] SKIP   {path.name} (empty after extraction)")
            continue

        if book_json.get("quality_status") == "quarantine":
            quarantined += 1
            report_entry = {
                "kind": "quality",
                **book_json,
                "quality_status": book_json.get("quality_status"),
                "quality_reasons": book_json.get("quality_reasons", []),
                "quality_warnings": book_json.get("quality_warnings", []),
                "quality_metrics": book_json.get("quality_metrics", {}),
                "expected_language": book_json.get("quality_expected_language"),
                "expected_arabic": book_json.get("quality_expected_arabic"),
                "detected_script": book_json.get("quality_detected_script"),
            }
            write_jsonl(QUALITY_REPORT_PATH, report_entry)
            log(
                f"[{i:04d}] QC     {path.name} "
                f"status={book_json.get('quality_status')} review={book_json.get('review_status')} "
                f"reasons={','.join(book_json.get('quality_reasons', [])) or '-'}"
            )
        elif book_json.get("quality_status") == "warn":
            quality_warned += 1
            log(
                f"[{i:04d}] WARN   {path.name} "
                f"reasons={','.join(book_json.get('quality_warnings', [])) or '-'}"
            )

        if duplicate_of and not args.keep_duplicates:
            skipped_duplicates += 1
            duplicate_hits += 1
            report_entry = {
                "source_path": book_json["source_path"],
                "source_relpath": book_json["source_relpath"],
                "source_ext": book_json.get("source_ext", ""),
                "source_type": book_json.get("source_type", "unknown"),
                "document_type": book_json.get("document_type", "book"),
                "conversion_status": book_json.get("conversion_status", "unknown"),
                "duplicate_of": duplicate_of.get("json_path"),
                "duplicate_title": duplicate_of.get("title"),
                "score": dup_info["score"] if dup_info else 0.0,
                "reason": dup_info["reason"] if dup_info else "duplicate",
                "content_hash": book_json["content_hash"],
                "language": book_json["language"],
                "title": book_json["title"],
                "extractor": extractor,
                "quality_status": book_json.get("quality_status", "ok"),
                "quality_reasons": book_json.get("quality_reasons", []),
                "review_status": book_json.get("review_status", "approved_auto"),
                "review_required": bool(book_json.get("review_required", False)),
            }
            write_duplicate_report(DUPLICATE_REPORT_PATH, report_entry)
            log(
                f"[{i:04d}] DUP    {path.name} -> {duplicate_of.get('json_path')} "
                f"score={report_entry['score']:.3f}"
            )
            continue

        outpath = OUTPUT_DIR / book_json["json_path"]
        outpath.parent.mkdir(parents=True, exist_ok=True)
        with outpath.open("w", encoding="utf-8") as f:
            json.dump(book_json, f, ensure_ascii=False, indent=2)

        index_record = {
            "book_id": book_json["book_id"],
            "filename": book_json["filename"],
            "language": book_json["language"],
            "title": book_json["title"],
            "total_pages": book_json["total_pages"],
            "size_bytes": book_json["size_bytes"],
            "json_path": book_json["json_path"],
            "source_root": book_json["source_root"],
            "source_path": book_json["source_path"],
            "source_relpath": book_json["source_relpath"],
            "source_ext": book_json.get("source_ext", ""),
            "source_type": book_json.get("source_type", "unknown"),
            "document_type": book_json.get("document_type", "book"),
            "conversion_status": book_json.get("conversion_status", "unknown"),
            "source_hash": book_json["source_hash"],
            "content_hash": book_json["content_hash"],
            "quality_status": book_json.get("quality_status", "ok"),
            "quality_reasons": book_json.get("quality_reasons", []),
            "quality_warnings": book_json.get("quality_warnings", []),
            "review_status": book_json.get("review_status", "approved_auto"),
            "review_required": bool(book_json.get("review_required", False)),
            "review_route": book_json.get("review_route", "auto"),
            "reviewed_by": book_json.get("reviewed_by", ""),
            "reviewed_at": book_json.get("reviewed_at", ""),
            "review_note": book_json.get("review_note", ""),
            "ingest_ready": bool(book_json.get("ingest_ready", True)),
        }

        update_index(index, index_record)
        signature = build_signature([p["content"] for p in book_json["pages"]], book_json["language"], book_json["size_bytes"])
        catalog = update_content_index(catalog, index_record, signature)
        imported += 1

        save_json(INDEX_PATH, index)
        save_json(CONTENT_INDEX_PATH, {"total_files": len(catalog), "entries": catalog})

        log(
            f"[{i:04d}] OK     {path.name} -> {book_json['json_path']} "
            f"pages={book_json['total_pages']} extractor={extractor}"
        )

    save_json(INDEX_PATH, index)
    save_json(CONTENT_INDEX_PATH, {"total_files": len(catalog), "entries": catalog})

    log("")
    log("=== DONE ===")
    log(f"Processed          : {processed}")
    log(f"Imported unique    : {imported}")
    log(f"Quality quarantined : {quarantined}")
    log(f"Quality warnings    : {quality_warned}")
    log(f"Skipped duplicates  : {skipped_duplicates}")
    log(f"Unsupported/empty   : {unsupported}")
    log(f"Duplicate matches   : {duplicate_hits}")
    log(f"Index path          : {INDEX_PATH}")
    log(f"Content index path  : {CONTENT_INDEX_PATH}")
    log(f"Duplicate report    : {DUPLICATE_REPORT_PATH}")

    if not args.skip_canonicalize:
        rc = run_canonicalizer(prune_qdrant=not args.no_prune_qdrant)
        if rc != 0:
            log(f"WARNING canonicalizer returned non-zero exit code: {rc}")


if __name__ == "__main__":
    main()
